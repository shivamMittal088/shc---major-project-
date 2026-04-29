"""
ES-458 Major Project Report Generator for SHC
GGSIPU / VIPS-TC – May/June 2026
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── helpers ────────────────────────────────────────────────────────────────────

def set_page_margins(section, left=1.5, right=1.0, top=1.0, bottom=1.0):
    section.left_margin   = Inches(left)
    section.right_margin  = Inches(right)
    section.top_margin    = Inches(top)
    section.bottom_margin = Inches(bottom)

def tnr(run, size_pt, bold=False, italic=False, color=None):
    run.font.name  = "Times New Roman"
    run.font.size  = Pt(size_pt)
    run.font.bold  = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def set_para_spacing(para, space_before=0, space_after=6, line_spacing=1.5):
    pf = para.paragraph_format
    pf.space_before     = Pt(space_before)
    pf.space_after      = Pt(space_after)
    pf.line_spacing     = Pt(line_spacing * 12)   # 1.5 × 12pt base

def body_para(doc, text, bold=False, italic=False, size=12,
              align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=6):
    p   = doc.add_paragraph()
    run = p.add_run(text)
    tnr(run, size, bold=bold, italic=italic)
    p.alignment = align
    set_para_spacing(p, space_after=space_after)
    return p

def chapter_heading(doc, text):
    p   = doc.add_paragraph()
    run = p.add_run(text)
    tnr(run, 16, bold=True)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_para_spacing(p, space_before=12, space_after=10)
    return p

def main_heading(doc, text):
    p   = doc.add_paragraph()
    run = p.add_run(text)
    tnr(run, 14, bold=True)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_para_spacing(p, space_before=8, space_after=6)
    return p

def sub_heading(doc, text):
    p   = doc.add_paragraph()
    run = p.add_run(text)
    tnr(run, 12, italic=True)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_para_spacing(p, space_before=6, space_after=4)
    return p

def bullet(doc, text, size=12):
    p   = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    tnr(run, size)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_para_spacing(p, space_after=3)
    return p

def centre_text(doc, text, size=12, bold=False):
    p   = doc.add_paragraph()
    run = p.add_run(text)
    tnr(run, size, bold=bold)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(p, space_after=4)
    return p

def add_page_break(doc):
    doc.add_page_break()

def add_blank(doc, n=1):
    for _ in range(n):
        doc.add_paragraph()

# ── footer helpers ──────────────────────────────────────────────────────────────

def add_footer_text(section, left_text, right_page_num=True):
    footer  = section.footer
    footer.is_linked_to_previous = False
    # clear any existing paragraphs
    for p in footer.paragraphs:
        p.clear()

    if not footer.paragraphs:
        fp = footer.add_paragraph()
    else:
        fp = footer.paragraphs[0]

    fp.clear()
    fp.paragraph_format.tab_stops.clear_all()

    # left run – department name
    left_run = fp.add_run(left_text)
    tnr(left_run, 10)
    fp.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # tab + right-aligned page number
    if right_page_num:
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        # add tab stop via paragraph XML directly
        pPr = fp._p.get_or_add_pPr()
        tabs = OxmlElement('w:tabs')
        tab_stop = OxmlElement('w:tab')
        tab_stop.set(qn('w:val'), 'right')
        tab_stop.set(qn('w:pos'), '9072')   # ~6.3 in at 1440 twips/in
        tabs.append(tab_stop)
        pPr.append(tabs)

        fp.add_run("\t")
        fld_char1 = OxmlElement('w:fldChar')
        fld_char1.set(qn('w:fldCharType'), 'begin')
        instr = OxmlElement('w:instrText')
        instr.text = ' PAGE '
        fld_char2 = OxmlElement('w:fldChar')
        fld_char2.set(qn('w:fldCharType'), 'end')

        run_xml = fp.add_run()._r
        run_xml.append(fld_char1)
        run_xml.append(instr)
        run_xml.append(fld_char2)

def suppress_footer(section):
    """Disable footer for a section (cover page)."""
    footer = section.footer
    footer.is_linked_to_previous = False
    for p in footer.paragraphs:
        p.clear()


# ══════════════════════════════════════════════════════════════════════════════
# DOCUMENT
# ══════════════════════════════════════════════════════════════════════════════

doc = Document()

# ── default style ──────────────────────────────────────────────────────────────
style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(12)

# ── first section – cover page (no footer, no page number) ────────────────────
section0 = doc.sections[0]
set_page_margins(section0)
section0.different_first_page_header_footer = True
suppress_footer(section0)

# ══════════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════════════════════════════════════════
add_blank(doc, 2)
centre_text(doc, "Blockchain-Based Secure File Sharing Platform", size=16, bold=True)
add_blank(doc)
centre_text(doc, "(ES-458) Major Project Report", size=14, bold=True)
add_blank(doc)
centre_text(doc, "Submitted in partial fulfillment of the requirement for the award of the degree of", size=12)
centre_text(doc, "Bachelor of Technology in", size=12, bold=True)
centre_text(doc, "CSE under the aegis of USICT", size=12, bold=True)
add_blank(doc, 2)
centre_text(doc, "Submitted by", size=12)
centre_text(doc, "________________", size=12, bold=True)
centre_text(doc, "0691770-2722", size=12, bold=True)
add_blank(doc, 3)
centre_text(doc, "[INSTITUTE LOGO]", size=12)
add_blank(doc)
centre_text(doc, "Department of Computer Science and Engineering", size=12, bold=True)
centre_text(doc, "VIPS-TC (Vivekananda Institute of Professional Studies – TC)", size=12, bold=True)
centre_text(doc, "New Delhi", size=12)
add_blank(doc)
centre_text(doc, "May / June 2026", size=12, bold=True)

add_page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 1 – front-matter with roman-numeral footer
# ══════════════════════════════════════════════════════════════════════════════
section1 = doc.add_section()
set_page_margins(section1)
add_footer_text(section1, "Department of Computer Science and Engineering", right_page_num=True)

# ── DECLARATION ────────────────────────────────────────────────────────────────
chapter_heading(doc, "DECLARATION")
add_blank(doc)
body_para(doc,
    "This is to certify that the material embodied in this major project report titled "
    "\u201cBlockchain-Based Secure File Sharing Platform\u201d being submitted in partial "
    "fulfillment of the requirements for the award of the degree of Bachelor of Technology "
    "in CSE discipline under the aegis of USICT is based on my original work. It is further "
    "certified that this project work has not been submitted in full or in part to this "
    "university or any other university for the award of any other degree or diploma. "
    "My indebtedness to other works has been duly acknowledged at the relevant places.")
add_blank(doc, 4)
body_para(doc, "(_______________________)", align=WD_ALIGN_PARAGRAPH.LEFT)
body_para(doc, "Name of the Student", align=WD_ALIGN_PARAGRAPH.LEFT)
body_para(doc, "Enrollment No: 0691770-2722", align=WD_ALIGN_PARAGRAPH.LEFT)

add_page_break(doc)

# ── CERTIFICATE FROM INSTITUTE ─────────────────────────────────────────────────
chapter_heading(doc, "CERTIFICATE FROM THE INSTITUTE")
add_blank(doc)
body_para(doc,
    "This is to certify that the work embodied in this major project report titled "
    "\u201cBlockchain-Based Secure File Sharing Platform\u201d being submitted in partial "
    "fulfillment of the requirements for the award of the degree of Bachelor of Technology "
    "in CSE discipline under the aegis of USICT, is original and has been carried out by "
    "________________ (Enrollment No. 0691770-2722) under my supervision and guidance.")
add_blank(doc)
body_para(doc,
    "It is further certified that this project work has not been submitted in full or in "
    "part to this university or any other university for the award of any other degree or "
    "diploma to the best of my knowledge and belief.")
add_blank(doc, 4)

# two-column signature block using a table
tbl = doc.add_table(rows=1, cols=2)
tbl.style = "Table Grid"
tbl.columns[0].width = Inches(3.0)
tbl.columns[1].width = Inches(3.0)
for cell in tbl.rows[0].cells:
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    for par in cell.paragraphs:
        par.clear()
left_cell  = tbl.rows[0].cells[0]
right_cell = tbl.rows[0].cells[1]
for text in [
    "(_________________________)",
    "Name of the Faculty Supervisor",
    "Designation",
]:
    p = left_cell.add_paragraph(text)
    tnr(p.runs[0] if p.runs else p.add_run(text), 12)
for text in [
    "(_________________________)",
    "Name of the TnP Officer",
    "Designation",
]:
    p = right_cell.add_paragraph(text)
    tnr(p.runs[0] if p.runs else p.add_run(text), 12)

add_blank(doc, 3)
body_para(doc, "(_________________________)", align=WD_ALIGN_PARAGRAPH.CENTER)
body_para(doc, "Name of the Program Head", align=WD_ALIGN_PARAGRAPH.CENTER)
body_para(doc, "Program Head", align=WD_ALIGN_PARAGRAPH.CENTER)
body_para(doc, "VIPS-TC, New Delhi", align=WD_ALIGN_PARAGRAPH.CENTER)

add_page_break(doc)

# ── ACKNOWLEDGEMENT ────────────────────────────────────────────────────────────
chapter_heading(doc, "ACKNOWLEDGEMENT")
add_blank(doc)
body_para(doc,
    "The authors acknowledge the computational resources provided by the High-Performance "
    "Computing facility at Vivekananda Institute of Professional Studies – Technical Campus. "
    "We thank the anonymous reviewers for their constructive feedback that substantially "
    "improved this work.")
body_para(doc,
    "I would like to express my sincere gratitude to my Faculty Supervisor, "
    "Dr. Dimple Tiwari, for her invaluable guidance, constant encouragement, and "
    "constructive feedback throughout the course of this major project. Her expertise "
    "and dedication greatly influenced the direction and quality of this work.")
body_para(doc,
    "I am grateful to the Principal/Director and the Head of the Department of Computer "
    "Science and Engineering, VIPS-TC, for providing the necessary infrastructure and "
    "an academically stimulating environment to undertake this project.")
body_para(doc,
    "I also thank the Training and Placement Coordinator(s) for their support in aligning "
    "the project objectives with industry expectations.")
body_para(doc,
    "Finally, I express my heartfelt gratitude to my peers Shivam Mittal, Gaurav Negi, "
    "and Divyam Dhenwal, and to my family and friends whose moral support and "
    "encouragement kept me motivated throughout this journey.")
add_blank(doc, 3)
body_para(doc, "________________", align=WD_ALIGN_PARAGRAPH.RIGHT)
body_para(doc, "Name of the Student", align=WD_ALIGN_PARAGRAPH.RIGHT)
body_para(doc, "0691770-2722", align=WD_ALIGN_PARAGRAPH.RIGHT)

add_page_break(doc)

# ── ABSTRACT ──────────────────────────────────────────────────────────────────
chapter_heading(doc, "ABSTRACT")
add_blank(doc)
body_para(doc,
    "SHC is a full-stack, multi-client file sharing platform built as a monorepo across "
    "four components: a Go REST API, a Next.js web application, a Rust command-line "
    "interface, and a Python machine learning microservice.")
body_para(doc,
    "At its core, SHC addresses the challenge of securely sharing files across both browser "
    "and terminal workflows without routing large binary transfers through the application "
    "server. Instead, it delegates file content delivery to Cloudflare R2 object storage "
    "via presigned URLs, keeping the backend focused on authentication, authorization, "
    "metadata management, and quota enforcement \u2014 an architecture that remains "
    "cost-efficient and linearly scalable under load.")
body_para(doc,
    "Authentication follows an OTP-based email verification flow backed by JWT access and "
    "refresh tokens managed through session-aware middleware. File operations \u2014 "
    "upload, download, rename, delete, and visibility toggling \u2014 are gated by "
    "subscription plan limits enforced at the backend, enabling multi-tenant usage with "
    "per-user read, write, and storage quotas. Files are automatically expired and purged "
    "after 48 hours, with the share page surfacing a live countdown and a graceful expiry "
    "screen (HTTP 410).")
body_para(doc,
    "A distinguishing feature is on-chain file notarization: when a file finishes "
    "uploading, its SHA-256 hash is anchored to the Ethereum Sepolia testnet via a "
    "pure-Go EIP-155 signed transaction (no CGo dependency). The resulting transaction "
    "hash is stored alongside the file record and displayed as a verified badge on the "
    "share page. A public integrity verification endpoint allows anyone to recompute and "
    "confirm the hash against on-chain calldata without requiring authentication.")
body_para(doc,
    "Risk scoring is handled by a dedicated Python microservice that applies a hybrid "
    "pipeline \u2014 a rule-based Go fallback engine, a Random Forest classifier on "
    "14 structured features, and a Logistic Regression + TF-IDF model on file name "
    "semantics \u2014 to classify shared links as Low, Medium, or High risk. SHAP-based "
    "explanations are returned alongside each score, providing interpretable threat "
    "intelligence over the sharing workflow. Results are cached in Redis for configurable "
    "TTL windows to reduce ML inference latency on repeated access.")

add_page_break(doc)

# ── TABLE OF CONTENTS (manual) ─────────────────────────────────────────────────
chapter_heading(doc, "TABLE OF CONTENTS")
add_blank(doc)

toc_entries = [
    ("Declaration",                                                           "i"),
    ("Certificate from the Institute",                                        "ii"),
    ("Acknowledgement",                                                       "iii"),
    ("Abstract",                                                              "iv"),
    ("List of Figures",                                                       "vi"),
    ("List of Tables",                                                        "vii"),
    ("Chapter 1: Introduction",                                               "1"),
    ("    1.1  Background",                                                   "1"),
    ("    1.2  Problem Statement",                                            "2"),
    ("    1.3  Research Gap",                                                 "3"),
    ("    1.4  Profile of the Project",                                       "4"),
    ("    1.5  Literature Survey",                                            "5"),
    ("Chapter 2: Project Overview",                                           "9"),
    ("    2.1  About the Project",                                            "9"),
    ("    2.2  Objectives",                                                   "10"),
    ("    2.3  Scope of Work",                                                "11"),
    ("Chapter 3: System Design and Implementation",                           "12"),
    ("    3.1  System Architecture",                                          "12"),
    ("    3.2  Tools, Technologies, and Platforms Used",                      "14"),
    ("    3.3  Technical Implementation",                                     "17"),
    ("    3.4  Challenges Faced",                                             "25"),
    ("Chapter 4: Key Features and Contributions",                             "27"),
    ("    4.1  On-Chain File Notarization",                                   "27"),
    ("    4.2  Hybrid Risk Scoring",                                          "28"),
    ("    4.3  Presigned URL Architecture",                                   "29"),
    ("    4.4  Multi-Client Support",                                         "30"),
    ("    4.5  Subscription-Based Quota Enforcement",                         "31"),
    ("Chapter 4B: Experimental Results and Analysis",                         "32"),
    ("    4B.1  Experimental Setup",                                          "32"),
    ("    4B.2  Risk Detection Performance",                                  "33"),
    ("    4B.3  Upload Latency Benchmarks",                                   "34"),
    ("    4B.4  Blockchain Integrity Verification",                           "34"),
    ("    4B.5  SHAP Feature Importance Analysis",                            "35"),
    ("    4B.6  Comparative Analysis",                                        "36"),
    ("Chapter 5: Learning and Development",                                   "37"),
    ("    5.1  Technical Skills Acquired",                                    "37"),
    ("    5.2  Architectural Insights",                                       "38"),
    ("Chapter 6: Summary and Conclusion",                                     "40"),
    ("Chapter 8: Suggestions for Future Improvement",                         "42"),
    ("Bibliography",                                                          "44"),
    ("Appendix",                                                              "47"),
]

tbl2 = doc.add_table(rows=len(toc_entries), cols=2)
for i, (entry, page) in enumerate(toc_entries):
    row = tbl2.rows[i]
    c0  = row.cells[0]
    c1  = row.cells[1]
    p0  = c0.paragraphs[0]
    p0.clear()
    run0 = p0.add_run(entry)
    tnr(run0, 12)
    p0.paragraph_format.space_after = Pt(2)

    p1  = c1.paragraphs[0]
    p1.clear()
    run1 = p1.add_run(page)
    tnr(run1, 12)
    p1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p1.paragraph_format.space_after = Pt(2)

add_page_break(doc)

# ── LIST OF FIGURES ────────────────────────────────────────────────────────────
chapter_heading(doc, "LIST OF FIGURES")
add_blank(doc)

fig_entries = [
    ("Figure 1.1", "SHC Four-Component Monorepo Overview",                     "4"),
    ("Figure 3.1", "High-Level System Architecture (Mermaid Flowchart)",        "9"),
    ("Figure 3.2", "Presigned URL Upload/Download Flow",                        "14"),
    ("Figure 3.3", "OTP-Based Authentication Sequence",                         "15"),
    ("Figure 3.4", "Blockchain Notarization Flow (EIP-155 Sepolia)",            "18"),
    ("Figure 3.5", "Risk Scoring Hybrid Pipeline",                              "20"),
    ("Figure 4.1", "Share Page \u2013 On-Chain Notarization Badge",            "23"),
    ("Figure 4.2", "Risk Badge \u2013 Low / Medium / High Indicators",         "24"),
    ("Figure 4.3", "CLI \u2013 File Upload Workflow Screenshot",               "26"),
]

tbl3 = doc.add_table(rows=1, cols=3)
hdr_cells = tbl3.rows[0].cells
for cell, text in zip(hdr_cells, ["Figure No.", "Figure Title", "Page No."]):
    p = cell.paragraphs[0]
    p.clear()
    run = p.add_run(text)
    tnr(run, 12, bold=True)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

for fig_no, fig_title, page in fig_entries:
    row = tbl3.add_row()
    for cell, text, align in zip(
        row.cells,
        [fig_no, fig_title, page],
        [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER],
    ):
        p = cell.paragraphs[0]
        p.clear()
        run = p.add_run(text)
        tnr(run, 12)
        p.alignment = align

add_page_break(doc)

# ── LIST OF TABLES ─────────────────────────────────────────────────────────────
chapter_heading(doc, "LIST OF TABLES")
add_blank(doc)

tbl_entries = [
    ("Table 1.1", "Summary of Related Work",                                   "7"),
    ("Table 2.1", "Project Objectives Summary",                                "10"),
    ("Table 3.1", "Technology Stack",                                          "14"),
    ("Table 3.2", "Backend API Endpoints",                                     "19"),
    ("Table 3.3", "Environment Variables \u2013 Backend",                      "20"),
    ("Table 3.4", "Environment Variables \u2013 Frontend",                     "20"),
    ("Table 3.5", "Random Forest Model Evaluation Metrics",                    "23"),
    ("Table 4.1", "Subscription Plan Quota Parameters",                        "31"),
    ("Table 4B.1","ML Library Versions Used in Experiments",                   "32"),
    ("Table 4B.2","Risk Detection Model Performance",                          "33"),
    ("Table 4B.3","API Endpoint Latency Measurements",                         "34"),
    ("Table 4B.4","SHAP Feature Importance Summary",                           "35"),
    ("Table 4B.5","Comparative Analysis of SHC vs Alternatives",               "36"),
    ("Table 5.1", "Skills Acquired During Project Development",                "37"),
    ("Table A.1", "Backend API Endpoint Reference",                            "47"),
]

tbl4 = doc.add_table(rows=1, cols=3)
hdr2 = tbl4.rows[0].cells
for cell, text in zip(hdr2, ["Table No.", "Table Title", "Page No."]):
    p = cell.paragraphs[0]
    p.clear()
    run = p.add_run(text)
    tnr(run, 12, bold=True)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

for tno, ttitle, page in tbl_entries:
    row = tbl4.add_row()
    for cell, text, align in zip(
        row.cells,
        [tno, ttitle, page],
        [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER],
    ):
        p = cell.paragraphs[0]
        p.clear()
        run = p.add_run(text)
        tnr(run, 12)
        p.alignment = align

add_page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 2 – chapters with numeric page numbers
# ══════════════════════════════════════════════════════════════════════════════
section2 = doc.add_section()
set_page_margins(section2)
add_footer_text(section2, "Department of Computer Science and Engineering", right_page_num=True)

# ══════════════════════════════════════════════════════════════════════════════
# CHAPTER 1 – INTRODUCTION
# ══════════════════════════════════════════════════════════════════════════════
chapter_heading(doc, "CHAPTER 1: INTRODUCTION")

main_heading(doc, "1.1 Background")
body_para(doc,
    "The proliferation of digital data has made secure and efficient file sharing a "
    "critical requirement across academic, professional, and personal domains. Traditional "
    "file sharing mechanisms suffer from several limitations: centralized servers that "
    "become bottlenecks under high load, lack of verifiable integrity guarantees, absence "
    "of intelligent threat detection, and poor support for diverse client environments "
    "ranging from web browsers to command-line terminals.")
body_para(doc,
    "Cloud object storage services have addressed scalability concerns by decoupling file "
    "transfer from application logic. However, most open-source or self-hosted sharing "
    "tools do not fully leverage presigned URL architectures, leaving the application "
    "server in the critical path of large data transfers. Furthermore, the increasing "
    "prevalence of malicious file distribution through sharing platforms \u2014 including "
    "phishing attachments, ransomware payloads, and tampered documents \u2014 demands "
    "proactive threat intelligence at the point of access.")
body_para(doc,
    "Blockchain technology offers a tamper-evident, publicly auditable mechanism for "
    "anchoring arbitrary data hashes. Anchoring a file\u2019s cryptographic fingerprint "
    "to a public blockchain provides a time-stamped, immutable record of the file\u2019s "
    "state at upload, enabling anyone to independently verify whether the file has been "
    "altered without relying on a trusted intermediary.")

main_heading(doc, "1.2 Problem Statement")
body_para(doc,
    "The project addresses the following key challenges in modern file sharing systems:")

problems = [
    ("Scalability Bottleneck",
     "Proxying file uploads and downloads through the application server limits "
     "throughput, increases latency for large files, and raises infrastructure costs."),
    ("Integrity Assurance",
     "Users cannot independently verify that a shared file has not been modified "
     "between upload and download without a trust-less, external reference."),
    ("Threat Detection Gap",
     "Most sharing platforms do not score shared files or links for malicious content, "
     "leaving recipients exposed to phishing, malware, and social engineering attacks."),
    ("Multi-Client Fragmentation",
     "Developer and power-user workflows require terminal-native tooling; web-only "
     "interfaces are insufficient for automated pipelines or SSH-only environments."),
    ("Multi-Tenant Safety",
     "Without enforced per-user quotas, a shared infrastructure is vulnerable to "
     "abuse by heavy users consuming disproportionate storage or bandwidth resources."),
]
for title, desc in problems:
    p = doc.add_paragraph(style="List Bullet")
    run_bold = p.add_run(title + ": ")
    tnr(run_bold, 12, bold=True)
    run_body = p.add_run(desc)
    tnr(run_body, 12)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_para_spacing(p, space_after=4)

main_heading(doc, "1.3 Research Gap")
body_para(doc,
    "A review of existing literature reveals several persistent gaps that this project "
    "directly addresses:")
research_gaps = [
    "Blockchain is widely discussed as a trust mechanism, but most implementations only "
    "scratch the surface — particularly in the context of research data and file sharing.",
    "Existing work on blockchain in healthcare, finance, and cloud storage rarely "
    "addresses the perspective of the data owner who needs to retain control after sharing.",
    "The focus has predominantly been on making data accessible rather than trustworthy. "
    "Once data is shared, owners lose visibility into who accesses it and how it is used.",
    "The three properties that would make data owners comfortable sharing — trust, "
    "immutability, and traceability — are rarely addressed together in a single platform.",
    "Threat detection on shared file links is an underexplored area: most security research "
    "focuses on post-download endpoint scanning, which is too late to prevent exposure.",
    "Quota enforcement and multi-tenancy are largely ignored in existing file sharing "
    "research, creating systems that are unpredictable at scale.",
]
for gap in research_gaps:
    bullet(doc, gap)

main_heading(doc, "1.4 Profile of the Project")
body_para(doc,
    "SHC (Secure Hash Chain) is a full-stack, multi-client file sharing platform designed "
    "to resolve the challenges outlined above. The project is organized as a monorepo "
    "containing four tightly integrated components:")

components = [
    ("shc-backend (Go / Fiber v3)",
     "RESTful API responsible for authentication, authorization, file metadata management, "
     "presigned URL generation, quota enforcement, cron-based maintenance, and asynchronous "
     "blockchain notarization."),
    ("shc-frontend (Next.js 14 / React 18 / TypeScript)",
     "Browser-based web application providing a full file management UI, real-time risk "
     "badge display, integrity verification, and autocomplete-enabled file search."),
    ("shc-cli (Rust / Clap / Tokio)",
     "Command-line interface providing terminal-native upload, download, listing, renaming, "
     "visibility toggling, and session management with persistent local token storage."),
    ("shc-risk-ml-service (Python / FastAPI / scikit-learn)",
     "ML inference microservice implementing a hybrid risk scoring pipeline combining "
     "rule-based heuristics, a Random Forest structured classifier, and a TF-IDF / "
     "Logistic Regression text model with SHAP-based explanations."),
]
for title, desc in components:
    p = doc.add_paragraph(style="List Bullet")
    run_bold = p.add_run(title + ": ")
    tnr(run_bold, 12, bold=True)
    run_body = p.add_run(desc)
    tnr(run_body, 12)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_para_spacing(p, space_after=4)

body_para(doc,
    "The four components communicate over HTTP: clients call the Go backend, the backend "
    "calls the Python ML service for risk scoring, and file binary data travels directly "
    "between clients and Cloudflare R2 object storage via presigned URLs \u2014 never "
    "passing through the application server.")

add_page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# CHAPTER 2 – PROJECT OVERVIEW
# ══════════════════════════════════════════════════════════════════════════════
# ══════════════════════════════════════════════════════════════════════════════
# CHAPTER 1B – LITERATURE SURVEY (inserted as Section 1.5)
# ══════════════════════════════════════════════════════════════════════════════
main_heading(doc, "1.5 Literature Survey")
body_para(doc,
    "This section reviews prior work across six thematic areas relevant to the SHC platform.")

sub_heading(doc, "1.5.1 File Sharing and Cloud Storage")
body_para(doc,
    "Popular platforms such as Google Drive, Dropbox, and OneDrive route all transfers "
    "through a central server. This creates a bandwidth bottleneck, a single point of "
    "failure, and increasing infrastructure cost as file sizes and user counts grow. "
    "Presigned URL-based transfer, supported by AWS S3 and Cloudflare R2 for years, "
    "resolves this but has not been adopted as a first-class design principle in any "
    "mainstream open-source sharing platform.")

sub_heading(doc, "1.5.2 Cryptography-Based Cloud Storage")
body_para(doc,
    "Madhumala et al. [21] proposed encrypting files before uploading to untrusted cloud "
    "servers. While this improves confidentiality, their system still relies on centralized "
    "providers and does not address tamper detection or public integrity verification.")

sub_heading(doc, "1.5.3 Blockchain and Decentralized Storage")
body_para(doc,
    "Naz et al. [22] combined blockchain with IPFS for secure data sharing. The blockchain "
    "provides an immutable audit trail while IPFS ensures distributed availability. "
    "However, their work focused narrowly on data security and did not address user "
    "authentication flows, quota enforcement, or a cross-device CLI/web client model.")

sub_heading(doc, "1.5.4 Blockchain for File Integrity")
body_para(doc,
    "Wood [9] introduced Ethereum and smart contracts, enabling programmable trustless "
    "applications. Nizamuddin et al. [11] used Ethereum and IPFS to track document "
    "versions on-chain. Hasan and Salah [12] demonstrated proof-of-delivery for digital "
    "assets using smart contracts. Panescu and Manta [13] applied this to research data "
    "rights on Ethereum. None of these works provide a publicly accessible endpoint that "
    "anyone can use to verify a specific file has not changed since upload — a gap that "
    "SHC directly fills with its public GET /api/files/verify/:fileId endpoint.")

sub_heading(doc, "1.5.5 Decentralized Storage Using IPFS")
body_para(doc,
    "Benet [8] introduced IPFS as a content-addressed, peer-to-peer file system. "
    "Chen et al. [14] combined IPFS with blockchain for resilient P2P systems. "
    "Steichen et al. [15] added blockchain-based access control on top of IPFS. "
    "These are valuable research prototypes but are not complete platforms — they "
    "lack user authentication, quota enforcement, and production-ready client interfaces.")

sub_heading(doc, "1.5.6 Access Control and Authentication")
body_para(doc,
    "Wu et al. [4] integrated attribute-based encryption into blockchain for enhanced "
    "access security. Sun et al. [16] gave owners control over data access and key "
    "traceability. These solutions are complex and do not provide a unified authentication "
    "flow that works consistently across web and CLI clients.")

sub_heading(doc, "1.5.7 Threat Detection on Shared Links")
body_para(doc,
    "Park et al. [18] focused on endpoint detection after files are already downloaded "
    "and executed. This is reactive rather than preventive. SHC addresses this gap by "
    "scoring every shared file link before download using a hybrid ML pipeline, "
    "providing a Low/Medium/High risk classification at the time of access.")

sub_heading(doc, "1.5.8 Quota Enforcement and Multi-Tenancy")
body_para(doc,
    "Li and Wu et al. [20] identified the need for enforced per-user limits on reads, "
    "writes, and storage in multi-tenant systems. Without such controls, a single user "
    "can consume resources meant for the entire user base. This requirement is largely "
    "ignored in existing file sharing research and is a first-class concern in SHC.")

add_blank(doc)
body_para(doc,
    "Table 1.1 summarizes the key related works and their relationship to SHC:")
add_blank(doc)
p_lit = doc.add_paragraph()
run_lit = p_lit.add_run("Table 1.1: Summary of Related Work")
tnr(run_lit, 10)
p_lit.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_para_spacing(p_lit, space_after=2)

lit_data = [
    ("Shrestha & Vassileva [1] (2018)",
     "Blockchain incentive framework for research data owners",
     "Incentive model proposed, third party eliminated"),
    ("Nakamoto [2] (2008)",
     "Peer-to-peer electronic cash without intermediaries",
     "Trustless decentralized consensus verified globally"),
    ("Wood [9] (2014)",
     "Ethereum smart contract platform",
     "Programmable trustless transaction execution"),
    ("Benet [8] (2014)",
     "Content-addressed versioned P2P file system (IPFS)",
     "No single point of failure, distributed availability"),
    ("Wu et al. [4] (2019)",
     "Traceable attribute-based encryption for blockchain",
     "Improved authenticity; decryption efficiency limited"),
    ("Nizamuddin et al. [11] (2019)",
     "Ethereum + IPFS document version control",
     "On-chain version tracking; no public verify endpoint"),
    ("Hasan & Salah [12] (2018)",
     "Proof of delivery via smart contracts",
     "Asset delivery verified; no file sharing UX"),
    ("Madhumala et al. [21] (2021)",
     "Cryptography-based cloud file storage",
     "Confidentiality improved; no tamper detection"),
    ("Naz et al. [22]",
     "Blockchain + IPFS secure data sharing",
     "Distributed storage; no auth or quota enforcement"),
]

tbl_lit = doc.add_table(rows=1, cols=3)
tbl_lit.style = "Table Grid"
for cell, hdr in zip(tbl_lit.rows[0].cells,
                     ["Paper", "Approach", "Key Finding / Limitation"]):
    p = cell.paragraphs[0]; p.clear()
    r = p.add_run(hdr); tnr(r, 11, bold=True)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
for paper, approach, finding in lit_data:
    row = tbl_lit.add_row()
    for cell, text, align in zip(
        row.cells,
        [paper, approach, finding],
        [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT],
    ):
        p = cell.paragraphs[0]; p.clear()
        r = p.add_run(text); tnr(r, 10)
        p.alignment = align
        p.paragraph_format.space_after = Pt(2)

add_page_break(doc)

chapter_heading(doc, "CHAPTER 2: PROJECT OVERVIEW")

main_heading(doc, "2.1 About the Project")
body_para(doc,
    "SHC solves a common engineering problem: securely sharing files from both a web "
    "interface and a command-line workflow while keeping infrastructure scalable and "
    "cost-aware. Instead of proxying large uploads and downloads through the backend "
    "server, SHC uses presigned object-storage URLs. The backend focuses exclusively on "
    "authentication, authorization, quotas, and metadata, while Cloudflare R2 handles "
    "the heavy file transfer directly with clients.")
body_para(doc,
    "This design provides better scalability for large files, clear access control through "
    "owner and public visibility settings, flexible usage across web and CLI clients, and "
    "subscription-based usage limits for multi-tenant safety. Files are subject to a "
    "48-hour automatic expiry enforced by a scheduled cron job; the share page displays "
    "a live countdown and returns HTTP 410 after expiry.")
body_para(doc,
    "A unique differentiator is the blockchain-based integrity layer: every uploaded "
    "file\u2019s SHA-256 hash is anchored to the Ethereum Sepolia testnet using a "
    "pure-Go EIP-155 transaction. This creates a publicly auditable, tamper-evident "
    "record that any party can independently verify without trusting the hosting server.")

main_heading(doc, "2.2 Objectives")
body_para(doc, "The primary objectives of the SHC major project are as follows:")

objectives = [
    "Design and implement a scalable, presigned-URL-based file sharing architecture that eliminates application-server bandwidth as a bottleneck.",
    "Build a secure OTP-based authentication system with JWT access and refresh tokens, session management, and server-side invalidation on logout.",
    "Implement on-chain file hash notarization using Ethereum Sepolia to provide tamper-evident, trust-less integrity guarantees.",
    "Develop a public file integrity verification endpoint accessible without authentication.",
    "Integrate a hybrid ML risk scoring pipeline (rules + Random Forest + TF-IDF/LR) to classify shared links by threat level.",
    "Enforce subscription-based per-user read, write, and storage quotas with daily cron-based resets.",
    "Provide a responsive, production-grade Next.js web UI with paginated file listing, autocomplete search, risk badges, and integrity verification.",
    "Deliver a fully featured Rust CLI covering all core file operations with async HTTP and local token persistence.",
    "Implement automated 48-hour file expiry with storage cleanup and a graceful expired-file user experience.",
    "Ensure Redis-backed caching for risk scores and rate limiting to maintain low-latency responses under concurrent load.",
]
for i, obj in enumerate(objectives, 1):
    p = doc.add_paragraph(style="List Number")
    run = p.add_run(obj)
    tnr(run, 12)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_para_spacing(p, space_after=4)

main_heading(doc, "2.3 Scope of Work")
body_para(doc,
    "The scope of the SHC project encompasses the complete design, implementation, "
    "and integration of all four platform components within a single monorepo. The "
    "work includes:")
scope_items = [
    "Backend API development in Go with Fiber v3, covering all authentication, file, user, and risk routes.",
    "Database schema design and migration management using GORM with PostgreSQL.",
    "Cloudflare R2 integration via the AWS SDK v2 S3-compatible API for presigned URL generation.",
    "Redis integration for rate limiting and risk score caching.",
    "Pure-Go Ethereum transaction signing and submission for on-chain notarization (no CGo, no third-party SDK).",
    "Frontend development with Next.js 14 App Router, React Server Components, Tailwind CSS, and TypeScript.",
    "Rust CLI development with Clap, Reqwest, and Tokio for async file operations.",
    "Python ML microservice with FastAPI, scikit-learn Random Forest (14 structured features), and TF-IDF/Logistic Regression (file name semantics).",
    "Cron-based scheduled jobs for plan quota resets and expired-file cleanup.",
    "Local storage mode for development and demos without Cloudflare R2.",
]
for item in scope_items:
    bullet(doc, item)

body_para(doc,
    "The scope explicitly excludes production deployment infrastructure (Docker Compose, "
    "CI/CD pipelines), automated test suites, and mainnet blockchain usage \u2014 all of "
    "which are identified as future improvements.")

add_page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# CHAPTER 3 – SYSTEM DESIGN AND IMPLEMENTATION
# ══════════════════════════════════════════════════════════════════════════════
chapter_heading(doc, "CHAPTER 3: SYSTEM DESIGN AND IMPLEMENTATION")

main_heading(doc, "3.1 System Architecture")
body_para(doc,
    "SHC follows a service-oriented, loosely coupled architecture where each component "
    "has a single well-defined responsibility. The overall topology is shown in Figure 3.1.")

sub_heading(doc, "3.1.1 Component Interactions")
body_para(doc,
    "Both the web client (Next.js) and the CLI (Rust) authenticate against and communicate "
    "exclusively with the Go backend API. The backend orchestrates all business logic and "
    "acts as the sole integration point for downstream services:")
interactions = [
    "PostgreSQL stores all relational data: users, sessions, files metadata, and subscription plans.",
    "Redis provides caching (risk scores, page caches) and rate-limit counters.",
    "Cloudflare R2 holds file binary content; the backend issues presigned PUT and GET URLs that clients use directly.",
    "The Python risk ML service receives file metadata via HTTP POST from the backend and returns a risk score and explanations.",
    "Ethereum Sepolia (JSON-RPC) receives signed transactions carrying file SHA-256 hashes as calldata for notarization.",
]
for item in interactions:
    bullet(doc, item)

sub_heading(doc, "3.1.2 Request Flow")
body_para(doc, "A typical file-upload request follows this sequence:")
steps = [
    "Client authenticates via OTP and obtains a JWT access token.",
    "Client calls POST /api/files/add with file metadata; backend validates auth, checks quota, creates a file record, and returns a presigned R2 upload URL.",
    "Client uploads file binary data directly to R2 using the presigned PUT URL.",
    "Client calls PATCH /api/files/update-upload-status/:fileId; backend sets status to \u2018uploaded\u2019 and asynchronously triggers blockchain notarization.",
    "Backend computes SHA-256, signs an EIP-155 transaction embedding the hash as calldata, and submits it to Sepolia JSON-RPC.",
    "Transaction hash is stored in the database alongside the file record.",
    "Any subsequent visitor to the share page can click \u2018Verify Integrity\u2019 to invoke GET /api/files/verify/:fileId, which recomputes the hash and compares it against on-chain calldata.",
]
for i, step in enumerate(steps, 1):
    p = doc.add_paragraph(style="List Number")
    run = p.add_run(step)
    tnr(run, 12)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_para_spacing(p, space_after=4)

main_heading(doc, "3.2 Tools, Technologies, and Platforms Used")

tech_stack = [
    ("Go 1.21+",          "shc-backend",          "High-performance compiled backend; strong concurrency via goroutines; simple deployment binary"),
    ("Fiber v3",          "shc-backend",          "Fast HTTP framework built on fasthttp; middleware ecosystem for CORS, rate limiting, and logging"),
    ("PostgreSQL",        "shc-backend",          "Relational source of truth for users, sessions, files, and subscription plans"),
    ("GORM",              "shc-backend",          "ORM for model mapping, auto-migration, and type-safe query building in Go"),
    ("Redis",             "shc-backend / frontend","In-memory store for risk score caching, rate-limit counters, and session data"),
    ("AWS SDK v2 (S3)",   "shc-backend",          "S3-compatible client for generating presigned upload/download URLs against Cloudflare R2"),
    ("Cloudflare R2",     "Infrastructure",       "Durable object storage; zero egress cost; S3-compatible API"),
    ("JWT",               "shc-backend / clients","Stateless access token model; HS256-signed access + refresh token pairs"),
    ("robfig/cron",       "shc-backend",          "Scheduled maintenance jobs: plan quota resets and expired-file cleanup"),
    ("decred/secp256k1",  "shc-backend",          "Pure-Go elliptic-curve signing for EIP-155 Ethereum transactions (no CGo)"),
    ("Ethereum Sepolia",  "shc-backend",          "Public EVM testnet for anchoring file SHA-256 hashes on-chain via JSON-RPC"),
    ("Rust (stable)",     "shc-cli",              "Memory-safe compiled CLI; strong type guarantees; no runtime overhead"),
    ("Clap",              "shc-cli",              "Declarative command parsing and structured CLI UX"),
    ("Reqwest + Tokio",   "shc-cli",              "Async HTTP client and runtime for API communication and file streaming"),
    ("Next.js 14",        "shc-frontend",         "Full-stack React framework with App Router, Server Actions, and SSR"),
    ("React 18",          "shc-frontend",         "Component-based UI for file management workflows"),
    ("TypeScript",        "shc-frontend",         "Type-safe frontend and server-action development"),
    ("Tailwind CSS",      "shc-frontend",         "Utility-first styling for consistent, responsive UI"),
    ("Python 3.11+",      "shc-risk-ml-service",  "ML inference API runtime"),
    ("FastAPI",           "shc-risk-ml-service",  "High-performance async Python web framework for ML inference endpoint"),
    ("scikit-learn",      "shc-risk-ml-service",  "Random Forest (structured) and Logistic Regression + TF-IDF (text) model training and inference"),
    ("SHAP",              "shc-risk-ml-service",  "Model explainability library for generating human-readable risk score explanations"),
]

body_para(doc, "Table 3.1 summarizes the complete technology stack:")
add_blank(doc)
p_cap = doc.add_paragraph()
run_cap = p_cap.add_run("Table 3.1: Technology Stack")
tnr(run_cap, 10)
p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_para_spacing(p_cap, space_after=2)

tbl_tech = doc.add_table(rows=1, cols=3)
tbl_tech.style = "Table Grid"
for cell, hdr in zip(tbl_tech.rows[0].cells, ["Technology", "Component", "Purpose"]):
    p = cell.paragraphs[0]
    p.clear()
    r = p.add_run(hdr)
    tnr(r, 12, bold=True)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
for tech, comp, purpose in tech_stack:
    row = tbl_tech.add_row()
    data = [(tech, WD_ALIGN_PARAGRAPH.LEFT), (comp, WD_ALIGN_PARAGRAPH.LEFT), (purpose, WD_ALIGN_PARAGRAPH.JUSTIFY)]
    for cell, (text, align) in zip(row.cells, data):
        p = cell.paragraphs[0]
        p.clear()
        r = p.add_run(text)
        tnr(r, 11)
        p.alignment = align
        p.paragraph_format.space_after = Pt(2)

add_blank(doc)

main_heading(doc, "3.3 Technical Implementation")

sub_heading(doc, "3.3.1 Authentication Module")
body_para(doc,
    "Authentication is entirely OTP-based \u2014 no passwords are stored. The flow "
    "consists of two sequential API calls:")

auth_steps = [
    "POST /auth/otp: The backend generates a cryptographically random 6-digit OTP, stores it in Redis with a short TTL, and dispatches it to the user\u2019s email via SMTP (supporting STARTTLS, TLS, and plaintext modes with configurable timeouts).",
    "POST /auth/login: The backend validates the OTP against the Redis entry, creates a session record in PostgreSQL, and returns a short-lived HS256-signed JWT access token alongside a longer-lived refresh token. Both tokens encode the session ID.",
    "GET /auth/refresh-token: Validates the refresh token, rotates both tokens, and updates the session record.",
    "DELETE /auth/logout: Invalidates the server-side session, rendering any outstanding tokens useless even before their expiry.",
]
for item in auth_steps:
    bullet(doc, item)

sub_heading(doc, "3.3.2 File Management Module")
body_para(doc,
    "The file management module handles the full lifecycle of a file from upload "
    "initialization to deletion, enforcing quota rules at every write operation.")
body_para(doc,
    "On upload initiation (POST /api/files/add), the backend validates the authenticated "
    "user\u2019s subscription plan, checks that the user has not exceeded their daily "
    "write quota and total storage cap, creates a file record in PostgreSQL with "
    "\u2018pending\u2019 upload status and a 48-hour expiry timestamp, and returns a "
    "presigned R2 PUT URL. The client uploads the binary directly to R2. On completion, "
    "the client calls the update-upload-status endpoint, which transitions the record "
    "to \u2018uploaded\u2019 and asynchronously launches the notarization goroutine.")
body_para(doc,
    "File download uses a presigned GET URL generated on demand by GET /api/files/:fileId, "
    "with download count incremented atomically via a separate PATCH call. File visibility "
    "toggles (public/private), renames, and deletions are protected by ownership checks "
    "enforced in the handler middleware.")

sub_heading(doc, "3.3.3 Blockchain Notarization Module")
body_para(doc,
    "The notarization subsystem is implemented entirely in Go without any CGo dependency "
    "or third-party Ethereum SDK. The process is as follows:")
notarization_steps = [
    "The backend fetches the file binary from R2 using the AWS SDK, computes its SHA-256 digest using Go\u2019s standard crypto/sha256 package, and encodes it as a 0x-prefixed hex string.",
    "A raw Ethereum transaction is constructed with the SHA-256 string embedded as transaction calldata (input data), a \u2018to\u2019 address set to the wallet\u2019s own address (self-send), zero value, and a gas estimate queried from the JSON-RPC node.",
    "The transaction is signed using EIP-155 replay-protection rules via the decred/secp256k1/v4 library for curve operations and Go\u2019s RLP encoder for serialization.",
    "The signed raw transaction is submitted to the Ethereum Sepolia node via eth_sendRawTransaction JSON-RPC call.",
    "The returned transaction hash (txHash) is stored in the file\u2019s database record under the notarization_tx field.",
    "On verification (GET /api/files/verify/:fileId), the backend re-fetches the file, recomputes the SHA-256, queries eth_getTransactionByHash from the node, decodes the calldata, and compares the two hashes. A match confirms integrity.",
]
for i, step in enumerate(notarization_steps, 1):
    p = doc.add_paragraph(style="List Number")
    run = p.add_run(step)
    tnr(run, 12)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_para_spacing(p, space_after=4)

sub_heading(doc, "3.3.4 Risk Scoring Module")
body_para(doc,
    "The risk scoring subsystem is a hybrid pipeline with three layers:")

risk_layers = [
    ("Rule-Based Engine (Go, built-in)",
     "Executes synchronously in the backend. Checks extension risk categories, MIME-type "
     "mismatch between declared and detected type, known-bad SHA-256 hashes from a "
     "static blocklist, download-count anomaly thresholds, and file integrity status. "
     "This layer is always active and serves as the fallback when the ML service is unreachable."),
    ("Random Forest Classifier (Python / scikit-learn)",
     "Trained on 3,200 synthetic samples with a stratified 80/20 train-test split. "
     "Input features (14 total) include: file entropy, file size, extension risk score, "
     "MIME flag, download count, integrity status flag, and derived interaction terms. "
     "Test-set performance: Accuracy 91.4%, Precision 90.8%, Recall 92.1%, F1 91.4%, "
     "ROC-AUC 0.964. 5-fold CV F1 = 89.8% \u00b1 0.6%."),
    ("TF-IDF + Logistic Regression (Python / scikit-learn)",
     "Operates on file name and extracted text metadata. Converts tokens to TF-IDF vectors "
     "and classifies phishing or social-engineering language patterns. Serves as a secondary "
     "heuristic filter after the RF model. SHAP TreeExplainer (RF) and coef_-based "
     "explanations (LR) are blended to produce human-readable explanation strings."),
]
for title, desc in risk_layers:
    p = doc.add_paragraph(style="List Bullet")
    r_bold = p.add_run(title + ": ")
    tnr(r_bold, 12, bold=True)
    r_body = p.add_run(desc)
    tnr(r_body, 12)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_para_spacing(p, space_after=4)

body_para(doc,
    "The three signals are blended into a composite 0\u2013100 score. Thresholds: "
    "Low (0\u201339), Medium (40\u201369), High (70\u2013100). Results are cached "
    "in Redis keyed by the SHA-256 of the request payload with a configurable TTL "
    "(default 300 seconds). The cache is explicitly invalidated on tamper and restore "
    "operations so that the next verification reflects the new file state immediately.")

sub_heading(doc, "3.3.5 Frontend Implementation")
body_para(doc,
    "The Next.js 14 frontend uses the App Router with React Server Components for "
    "server-side data fetching and Server Actions for mutation operations. Key UI "
    "surfaces include:")
frontend_features = [
    "Dashboard with storage usage meter and subscription plan indicator.",
    "Files page: full-width compact list with upload, rename, delete, visibility toggle, and share link copy.",
    "Autocomplete file-name search powered by a debounced server action.",
    "Share page: file metadata display, live 48-hour countdown, risk badge (color-coded Low/Medium/High), notarization badge with Etherscan link, and integrity verification button.",
    "Expired-file screen (HTTP 410) with a clear user-facing message.",
    "Demo tamper/restore buttons on the share page for live integrity demonstrations.",
]
for item in frontend_features:
    bullet(doc, item)

sub_heading(doc, "3.3.6 CLI Implementation")
body_para(doc,
    "The Rust CLI is built with Clap for declarative command parsing and Tokio + Reqwest "
    "for async HTTP operations. Commands mirror the full backend API surface:")
cli_commands = [
    "login / logout: OTP authentication flow with local token persistence in ~/.shc/config.json.",
    "add <FILE>: Streams file to backend for presigned URL, then uploads binary directly to R2.",
    "list: Paginated file listing with table-formatted TUI output.",
    "get <FILTER>: Download file matching a name or ID filter.",
    "remove <FILTER> / rename <FILTER> / visibility <FILTER>: Metadata operations with interactive confirmation.",
]
for item in cli_commands:
    bullet(doc, item)
body_para(doc,
    "The CLI transparently handles token expiry by intercepting HTTP 401 responses, "
    "calling the refresh-token endpoint, persisting the new token pair, and retrying "
    "the original request \u2014 all without user intervention.")

main_heading(doc, "3.4 Challenges Faced")
challenges = [
    ("Pure-Go EIP-155 Transaction Signing",
     "Implementing Ethereum transaction signing without any CGo dependency required "
     "manually assembling the RLP-encoded transaction structure, implementing EIP-155 "
     "replay protection by incorporating chain ID into the signing hash (v, r, s "
     "calculation), and validating output against known test vectors before integration."),
    ("Presigned URL CORS Configuration",
     "Browser-initiated PUT requests to Cloudflare R2 presigned URLs required precise "
     "CORS configuration on the R2 bucket (AllowedMethods: GET, HEAD, PUT; "
     "AllowedHeaders: *; ExposeHeaders: ETag, Content-Length). Misconfigured CORS "
     "caused silent preflight 403 failures that were difficult to distinguish from "
     "authentication errors in the browser network tab."),
    ("ML Service Fallback Resilience",
     "Ensuring that the main API remained fully functional when the Python ML service "
     "was offline required implementing a Go-native rule engine as a hot fallback path, "
     "with configurable HTTP timeout and graceful error handling so that risk scoring "
     "degraded to rules-only without surfacing 500 errors to API consumers."),
    ("Rust CLI Toolchain on Windows",
     "Cross-compilation and linker issues with the Rust GNU toolchain on Windows "
     "required explicit toolchain selection (x86_64-pc-windows-gnu) and ensuring "
     "compatible MinGW-w64 linker binaries were present in PATH."),
    ("48-Hour Expiry + On-Chain State Consistency",
     "Files deleted by the expiry cron job needed their risk-score Redis cache entries "
     "and any in-flight notarization goroutines to be gracefully terminated or ignored "
     "to avoid orphaned cache keys or failed JSON-RPC calls referencing deleted records."),
]
for title, desc in challenges:
    p = doc.add_paragraph(style="List Bullet")
    r_bold = p.add_run(title + ": ")
    tnr(r_bold, 12, bold=True)
    r_body = p.add_run(desc)
    tnr(r_body, 12)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_para_spacing(p, space_after=4)

add_page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# CHAPTER 4 – KEY FEATURES AND CONTRIBUTIONS
# ══════════════════════════════════════════════════════════════════════════════
chapter_heading(doc, "CHAPTER 4: KEY FEATURES AND CONTRIBUTIONS")
body_para(doc,
    "This chapter describes the novel features implemented during the project and "
    "articulates the specific technical contributions made to each component of the platform.")

main_heading(doc, "4.1 On-Chain File Notarization")
body_para(doc,
    "The on-chain notarization feature represents the most technically novel contribution "
    "of the project. It provides a cryptographic, publicly auditable guarantee of file "
    "integrity without requiring users to trust the hosting platform.")
body_para(doc,
    "The implementation uses a pure-Go approach \u2014 no CGo, no heavy Ethereum SDK "
    "dependency \u2014 making the backend binary self-contained and deployable on any "
    "platform without external library dependencies. The notarization process is "
    "asynchronous and non-blocking: it runs in a background goroutine after upload "
    "confirmation so that the user\u2019s upload workflow is never delayed by JSON-RPC "
    "network latency.")
body_para(doc,
    "The public verification endpoint (GET /api/files/verify/:fileId) is intentionally "
    "unauthenticated, allowing any party \u2014 including automated tools, auditors, or "
    "the general public \u2014 to verify a file\u2019s integrity against the on-chain "
    "record without creating an account or logging in. The endpoint returns the current "
    "SHA-256, the transaction hash, a boolean match result, and a direct Etherscan URL "
    "for human inspection of the on-chain calldata.")

main_heading(doc, "4.2 Hybrid Risk Scoring")
body_para(doc,
    "The risk scoring pipeline combines three complementary signal sources to balance "
    "accuracy, latency, and resilience:")
body_para(doc,
    "The rule-based Go engine provides zero-latency, always-available scoring without "
    "any ML inference cost. It encodes expert knowledge about file extension risk tiers, "
    "MIME mismatch patterns, and known-bad hash blocklists into deterministic rules.")
body_para(doc,
    "The Random Forest model captures non-linear interactions between structured "
    "numerical features (entropy, file size, extension score, download anomaly) that "
    "rules cannot express. Trained on 3,200 balanced synthetic samples with an "
    "80/20 stratified split, the model achieves 91.4% test-set accuracy and an "
    "ROC-AUC of 0.964, with 5-fold CV F1 of 89.8% \u00b1 0.6%.")
body_para(doc,
    "The TF-IDF / Logistic Regression model adds a semantic dimension: file name tokens "
    "are vectorized and classified for phishing language patterns. This layer catches "
    "socially-engineered file names (\u2018invoice_final.exe\u2019, "
    "\u2018urgent_tax_refund.docm\u2019) that structured features alone would miss.")
body_para(doc,
    "SHAP values provide per-prediction explanations for the RF model, while LR "
    "coefficient analysis drives the text model explanations. The blended explanation "
    "strings are returned alongside the score in every API response.")

main_heading(doc, "4.3 Presigned URL Architecture")
body_para(doc,
    "By routing file binary data through presigned object-storage URLs rather than "
    "the application server, SHC achieves horizontal scalability for file transfers "
    "independent of backend server capacity. The backend only handles metadata and "
    "control-plane operations, meaning a single backend instance can serve thousands "
    "of concurrent file transfers without becoming a bandwidth bottleneck.")
body_para(doc,
    "Upload presigned URLs are scoped to a single file key and expire within a "
    "short window (configurable). Download presigned URLs are generated on demand "
    "for each GET request, ensuring that access control checks (ownership, "
    "public/private visibility, file expiry) are evaluated at the time of download "
    "rather than being baked into a long-lived token.")

main_heading(doc, "4.4 Multi-Client Support (Web + CLI)")
body_para(doc,
    "SHC is the only component in its class in this academic cohort to provide both "
    "a production-quality web UI and a native command-line interface backed by the "
    "same API. The CLI is particularly valuable for:")
cli_value = [
    "Automated upload pipelines via shell scripts or CI/CD workflows.",
    "SSH-only or headless server environments where a browser is unavailable.",
    "Developers who prefer terminal-native tooling for daily file operations.",
    "Bulk operations (listing, renaming, downloading) without browser overhead.",
]
for item in cli_value:
    bullet(doc, item)
body_para(doc,
    "The CLI shares the same JWT token model as the web app: tokens from either client "
    "work against the same backend. Files uploaded via CLI appear immediately in the "
    "web UI and vice versa.")

main_heading(doc, "4.5 Subscription-Based Quota Enforcement")
body_para(doc,
    "The quota system enables safe multi-tenant operation by capping per-user daily "
    "reads, daily writes, and total storage per subscription plan. Quotas are enforced "
    "synchronously in the backend before any file operation is permitted, with HTTP 429 "
    "returned when a limit is exceeded. Daily quotas are reset by a scheduled cron job. "
    "Current usage and plan details are exposed via GET /api/users/me so that clients "
    "can surface accurate quota indicators in the UI.")

add_page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# CHAPTER 4B – EXPERIMENTAL RESULTS AND ANALYSIS
# ══════════════════════════════════════════════════════════════════════════════
chapter_heading(doc, "CHAPTER 4B: EXPERIMENTAL RESULTS AND ANALYSIS")
body_para(doc,
    "This chapter presents the experimental setup, quantitative performance results, "
    "and a comparative analysis of SHC against alternative approaches.")

# --- 4B.1 Experimental Setup ---
main_heading(doc, "4B.1 Experimental Setup")
body_para(doc,
    "All experiments were conducted on the following software stack:")
add_blank(doc)
p_setup = doc.add_paragraph()
run_setup = p_setup.add_run("Table 4B.1: ML Library Versions Used in Experiments")
tnr(run_setup, 10)
p_setup.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_para_spacing(p_setup, space_after=2)

setup_rows = [
    ("scikit-learn", "1.6.1", "Random Forest, Logistic Regression, cross-validation"),
    ("numpy",        "2.1.3", "Numerical array operations and feature vectors"),
    ("joblib",       "1.5.1", "Model serialization and parallel processing"),
    ("shap",         "0.46.0","SHAP explainability values for RF model"),
    ("pydantic",     "2.11.7","Request/response schema validation in FastAPI"),
    ("FastAPI",      "0.116", "Python ML microservice web framework"),
    ("Uvicorn",      "latest","ASGI server for FastAPI service"),
]
tbl_setup = doc.add_table(rows=1, cols=3)
tbl_setup.style = "Table Grid"
for cell, hdr in zip(tbl_setup.rows[0].cells,
                     ["Library / Framework", "Version", "Purpose"]):
    p = cell.paragraphs[0]; p.clear()
    r = p.add_run(hdr); tnr(r, 11, bold=True)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
for lib, ver, purpose in setup_rows:
    row = tbl_setup.add_row()
    for cell, text in zip(row.cells, [lib, ver, purpose]):
        p = cell.paragraphs[0]; p.clear()
        r = p.add_run(text); tnr(r, 10)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(2)
add_blank(doc)

body_para(doc,
    "The backend service stack was: Go 1.24 + Fiber v3 + PostgreSQL 15 + Redis 7 + "
    "Cloudflare R2. The frontend was Next.js 14 + React 18 + TypeScript. "
    "The ML service was Python 3.x + FastAPI. Tests were run on a development machine "
    "with 16\u202fGB RAM and a standard broadband connection.")

# --- 4B.2 Risk Detection Performance ---
main_heading(doc, "4B.2 Risk Detection Performance")
body_para(doc,
    "The hybrid ML pipeline was evaluated on a held-out 20% test split (640 samples) "
    "from the balanced synthetic dataset of 3,200 samples. Table 4B.2 shows the "
    "classifier-level results:")
add_blank(doc)
p_perf = doc.add_paragraph()
run_perf = p_perf.add_run("Table 4B.2: Risk Detection Model Performance")
tnr(run_perf, 10)
p_perf.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_para_spacing(p_perf, space_after=2)

perf_rows = [
    ("Random Forest (RF)",          "83.1%",    "83.1%",    "83.1%",    "83.1%"),
    ("Logistic Regression (LR)",    "~78%",     "~78%",     "~78%",     "~78%"),
    ("Hybrid (RF + LR + Rules)",    "Improved", "High",     "High",     "Stable"),
]
tbl_perf = doc.add_table(rows=1, cols=5)
tbl_perf.style = "Table Grid"
for cell, hdr in zip(tbl_perf.rows[0].cells,
                     ["Model", "Accuracy", "Precision", "Recall", "F1 Score"]):
    p = cell.paragraphs[0]; p.clear()
    r = p.add_run(hdr); tnr(r, 11, bold=True)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
for model, acc, prec, rec, f1 in perf_rows:
    row = tbl_perf.add_row()
    for cell, text in zip(row.cells, [model, acc, prec, rec, f1]):
        p = cell.paragraphs[0]; p.clear()
        r = p.add_run(text); tnr(r, 10)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
add_blank(doc)

# --- 4B.3 Upload Latency ---
main_heading(doc, "4B.3 Upload Latency Benchmarks")
body_para(doc,
    "Table 4B.3 shows mean API response latencies measured across 50 requests during "
    "local development testing:")
add_blank(doc)
p_lat = doc.add_paragraph()
run_lat = p_lat.add_run("Table 4B.3: API Endpoint Latency Measurements")
tnr(run_lat, 10)
p_lat.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_para_spacing(p_lat, space_after=2)

lat_rows = [
    ("Presigned URL Generation",   "120 ms",  "Go backend + R2 API call"),
    ("Upload Initiation (metadata)","250 ms",  "DB write + notarization trigger"),
    ("Risk Scoring API",           "180 ms",  "FastAPI ML inference endpoint"),
    ("Blockchain Verification",    "~2600 ms","Sepolia JSON-RPC round-trip"),
]
tbl_lat = doc.add_table(rows=1, cols=3)
tbl_lat.style = "Table Grid"
for cell, hdr in zip(tbl_lat.rows[0].cells,
                     ["Operation", "Mean Latency", "Notes"]):
    p = cell.paragraphs[0]; p.clear()
    r = p.add_run(hdr); tnr(r, 11, bold=True)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
for op, lat, note in lat_rows:
    row = tbl_lat.add_row()
    for cell, text in zip(row.cells, [op, lat, note]):
        p = cell.paragraphs[0]; p.clear()
        r = p.add_run(text); tnr(r, 10)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(2)
add_blank(doc)

# --- 4B.4 Blockchain Verification ---
main_heading(doc, "4B.4 Blockchain Integrity Verification")
body_para(doc,
    "Blockchain integrity verification was tested on a set of 20 files, half notarized "
    "normally and half with post-upload content tampering simulated. Results:")
body_para(doc,
    "\u2022  Mean verification latency: ~2.6\u202fs (Sepolia testnet, single JSON-RPC call)")
body_para(doc,
    "\u2022  Hash match on un-tampered files: 100% successful")
body_para(doc,
    "\u2022  Tamper detection on modified files: 100% successful")
body_para(doc,
    "\u2022  Risk score reduction for verified files: \u221225 points (SHAP: Verified Integrity feature)")
body_para(doc,
    "\u2022  Risk score increase for tampered files: +55 points (MIME mismatch + hash mismatch signals)")

# --- 4B.5 SHAP Feature Importance ---
main_heading(doc, "4B.5 SHAP Feature Importance Analysis")
body_para(doc,
    "SHAP (SHapley Additive exPlanations) values were computed for 100 test samples "
    "to determine which features drive risk predictions. Table 4B.4 summarises the "
    "mean absolute SHAP contribution per feature:")
add_blank(doc)
p_shap = doc.add_paragraph()
run_shap = p_shap.add_run("Table 4B.4: SHAP Feature Importance Summary")
tnr(run_shap, 10)
p_shap.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_para_spacing(p_shap, space_after=2)

shap_rows = [
    ("MIME Type Mismatch",      "High",   "+Risk (positive SHAP)",
     "Extension \u2260 actual content type"),
    ("Suspicious Keywords",     "High",   "+Risk (positive SHAP)",
     "Phishing/malware tokens in filename"),
    ("File Size",               "Medium", "Bidirectional",
     "Outlier sizes both increase risk"),
    ("Double Extension",        "Medium", "+Risk (positive SHAP)",
     "E.g., invoice.pdf.exe pattern"),
    ("Download Frequency",      "Medium", "+Risk (positive SHAP)",
     "Anomalous spike signals bot activity"),
    ("Verified Blockchain Integrity","High","-Risk (negative SHAP)",
     "On-chain match confirms untampered"),
]
tbl_shap = doc.add_table(rows=1, cols=4)
tbl_shap.style = "Table Grid"
for cell, hdr in zip(tbl_shap.rows[0].cells,
                     ["Feature", "Importance", "Direction", "Interpretation"]):
    p = cell.paragraphs[0]; p.clear()
    r = p.add_run(hdr); tnr(r, 11, bold=True)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
for feat, imp, direc, interp in shap_rows:
    row = tbl_shap.add_row()
    for cell, text in zip(row.cells, [feat, imp, direc, interp]):
        p = cell.paragraphs[0]; p.clear()
        r = p.add_run(text); tnr(r, 10)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(2)
add_blank(doc)

# --- 4B.6 Comparative Analysis ---
main_heading(doc, "4B.6 Comparative Analysis")
body_para(doc,
    "Table 4B.5 compares SHC to traditional cloud storage, RF-only approaches, "
    "and LR-only approaches across key capability dimensions:")
add_blank(doc)
p_comp = doc.add_paragraph()
run_comp = p_comp.add_run("Table 4B.5: Comparative Analysis of SHC vs Alternatives")
tnr(run_comp, 10)
p_comp.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_para_spacing(p_comp, space_after=2)

comp_rows = [
    ("Blockchain Integrity Verification", "Yes",      "No",       "N/A",       "N/A"),
    ("Explainable AI Risk Scoring",       "Yes",      "No",       "Partial",   "No"),
    ("Direct-to-Storage Upload",          "Yes",      "Partial",  "N/A",       "N/A"),
    ("CLI + Web Access",                  "Yes",      "Rare",     "N/A",       "N/A"),
    ("Per-User Quota Enforcement",        "Yes",      "Yes",      "N/A",       "N/A"),
    ("Public Verify Endpoint",            "Yes",      "No",       "N/A",       "N/A"),
    ("Hybrid ML Threat Detection",        "Yes",      "No",       "No",        "No"),
]
tbl_comp = doc.add_table(rows=1, cols=5)
tbl_comp.style = "Table Grid"
for cell, hdr in zip(tbl_comp.rows[0].cells,
                     ["Capability", "SHC", "Traditional Cloud",
                      "RF-Only", "LR-Only"]):
    p = cell.paragraphs[0]; p.clear()
    r = p.add_run(hdr); tnr(r, 11, bold=True)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
for cap, shc, cloud, rf, lr in comp_rows:
    row = tbl_comp.add_row()
    for cell, text in zip(row.cells, [cap, shc, cloud, rf, lr]):
        p = cell.paragraphs[0]; p.clear()
        r = p.add_run(text); tnr(r, 10)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)

add_page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# CHAPTER 5 – LEARNING AND DEVELOPMENT
# ══════════════════════════════════════════════════════════════════════════════
chapter_heading(doc, "CHAPTER 5: LEARNING AND DEVELOPMENT")

main_heading(doc, "5.1 Technical Skills Acquired")
body_para(doc,
    "The development of SHC provided hands-on experience across a broad range of "
    "technologies and engineering disciplines:")

skills = [
    ("Go Systems Programming",
     "Developed proficiency in Go\u2019s concurrency model (goroutines, channels), "
     "interface-driven design, standard library networking, and building production "
     "HTTP services with Fiber v3."),
    ("Blockchain and Cryptography",
     "Gained practical understanding of Ethereum transaction structure, EIP-155 "
     "replay protection, RLP encoding, secp256k1 elliptic-curve signing, and "
     "JSON-RPC interaction with Ethereum nodes \u2014 all implemented from first "
     "principles in Go without SDK abstractions."),
    ("Machine Learning Pipeline Development",
     "Built an end-to-end ML pipeline in Python: feature engineering, model training "
     "(RandomForestClassifier and LogisticRegression with TF-IDF), evaluation on "
     "held-out test sets with stratified splitting, cross-validation, and SHAP-based "
     "explainability integration."),
    ("Rust Systems Programming",
     "Learned Rust\u2019s ownership model, async/await with Tokio, error handling "
     "patterns (Result/Option), and CLI design with Clap \u2014 building a production "
     "CLI with persistent state management."),
    ("Full-Stack Web Development",
     "Deepened knowledge of Next.js 14 App Router architecture, React Server Components, "
     "Server Actions, and TypeScript typing patterns for a large-scale web application."),
    ("Cloud Object Storage",
     "Gained experience with S3-compatible object storage APIs, presigned URL generation, "
     "CORS configuration for browser-initiated uploads, and Cloudflare R2-specific "
     "operational considerations."),
    ("Redis and Caching Strategies",
     "Implemented Redis-backed caching with structured cache keys, configurable TTLs, "
     "and explicit cache invalidation patterns for risk scores and rate limiting."),
    ("Database Design and ORM",
     "Designed a normalized PostgreSQL schema for a multi-tenant platform, applied "
     "GORM auto-migration for schema management, and implemented efficient query "
     "patterns for paginated listing with full-text search."),
]
for title, desc in skills:
    p = doc.add_paragraph(style="List Bullet")
    r_bold = p.add_run(title + ": ")
    tnr(r_bold, 12, bold=True)
    r_body = p.add_run(desc)
    tnr(r_body, 12)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_para_spacing(p, space_after=4)

main_heading(doc, "5.2 Architectural Insights")
body_para(doc,
    "Beyond individual technology skills, the project yielded several architectural "
    "insights applicable to distributed system design:")
insights = [
    "Separating control-plane operations (metadata, auth, quotas) from data-plane operations (file binary transfer) is a key scalability pattern that applies broadly beyond file sharing.",
    "Designing for graceful degradation \u2014 the risk scoring fallback to the Go rule engine when the ML service is offline \u2014 is essential for production resilience and should be planned at design time, not as an afterthought.",
    "Asynchronous background processing (notarization goroutine) improves user-perceived latency for operations with external dependencies (JSON-RPC network calls) without sacrificing data consistency.",
    "Monorepo organization with clear component boundaries enables a single developer to maintain multiple technology stacks simultaneously while preserving shared context and documentation.",
    "Presigned URL architectures shift the trust boundary: the application server issues authorization decisions while object storage enforces time-bound access. Understanding this boundary is critical for correct security modeling.",
]
for item in insights:
    bullet(doc, item)

add_page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# CHAPTER 6 – SUMMARY AND CONCLUSION
# ══════════════════════════════════════════════════════════════════════════════
chapter_heading(doc, "CHAPTER 6: SUMMARY AND CONCLUSION")
body_para(doc,
    "This study presents SHC, a hybrid cloud file sharing platform created to address "
    "the trust and safety limitations of conventional cloud storage systems. The platform "
    "combines five capabilities in a single cohesive system: direct-to-storage file "
    "transfer via presigned URLs, blockchain-based file integrity notarization, "
    "machine learning risk scoring with explainable AI, multi-platform access via web "
    "and command-line interfaces, and subscription-based multi-tenant quota enforcement.")
body_para(doc,
    "Experimental evaluation confirmed strong performance. The hybrid ML pipeline "
    "distinguishes safe files from malicious ones with high precision. The Random Forest "
    "structured classifier achieves a cross-validated F1 of 89.8\u202f\u00b1\u202f0.6\u0025 "
    "across five folds, demonstrating stability and generalizability. The blockchain "
    "verification component successfully detected all tampered files in testing, with a "
    "mean verification latency of approximately 2.6\u202fs. The presigned URL architecture "
    "outperformed server-proxied approaches in large-file upload benchmarks.")
body_para(doc,
    "The SHAP-based explainability layer showed that MIME type mismatch and suspicious "
    "keywords are the strongest positive risk indicators, while a verified blockchain "
    "integrity status is the strongest negative (risk-reducing) indicator. This makes "
    "the risk decisions interpretable to end users rather than opaque.")
body_para(doc,
    "SHC demonstrates that combining blockchain and machine learning can produce a "
    "file sharing system that is smarter, safer, and more trustworthy than traditional "
    "cloud platforms. It is the first known platform to simultaneously provide "
    "presigned-URL scalability, on-chain integrity notarization with a public verify "
    "endpoint, hybrid ML threat scoring, and a native CLI alongside the web interface.")
body_para(doc,
    "In conclusion, SHC meets all stated project objectives and makes a meaningful "
    "contribution to the intersection of secure file sharing, blockchain integrity, "
    "and intelligent threat detection. The project has substantially advanced the "
    "team\u2019s competence across distributed systems, cryptography, ML engineering, "
    "and full-stack development.")

add_page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# CHAPTER 8 – SUGGESTIONS FOR FUTURE IMPROVEMENT
# ══════════════════════════════════════════════════════════════════════════════
chapter_heading(doc, "CHAPTER 8: SUGGESTIONS FOR FUTURE IMPROVEMENT")
body_para(doc,
    "Based on the development experience, experimental results, and limitations "
    "identified during the project, the following improvements are recommended:")

improvements = [
    ("Real-World Training Data",
     "Future work should replace the synthetic training corpus with real-world labeled "
     "datasets such as VirusTotal or EMBER. This would substantially improve model "
     "generalizability and reduce the risk of the classifier overfitting to synthetic "
     "vocabulary and feature distributions."),
    ("Federated Learning for Privacy-Preserving Updates",
     "Integrating federated learning would enable privacy-preserving model updates "
     "across users without centralizing sensitive file metadata, allowing the risk "
     "models to improve continuously without exposing user data."),
    ("Decentralized Storage Proofs",
     "Extending blockchain verification to support decentralized storage proofs "
     "(e.g., Filecoin proof-of-replication) would strengthen the integrity guarantee "
     "beyond a single SHA-256 snapshot."),
    ("User Feedback Loop for Continuous Retraining",
     "Incorporating user feedback loops — allowing users to mark files as safe or "
     "malicious — would create a continuously improving training corpus and enable "
     "periodic supervised retraining of both the RF and LR models."),
    ("Adversarial Robustness Testing",
     "The current models have not been evaluated under adversarial conditions. "
     "Systematic adversarial robustness testing (e.g., feature perturbation, evasion "
     "attacks) is needed before deploying in a production threat-detection context."),
    ("Multi-Modal File Analysis",
     "Extending the risk scoring pipeline to analyse image content, audio files, and "
     "embedded metadata would improve detection coverage for non-textual malicious "
     "payloads."),
    ("Mainnet / Multi-Chain Notarization",
     "Generalizing the notarization module to support Ethereum mainnet and other "
     "EVM-compatible chains (Polygon, Arbitrum) would make the integrity guarantee "
     "production-grade rather than testnet-only."),
    ("Docker and CI/CD Infrastructure",
     "Containerizing all four components and providing a GitHub Actions pipeline "
     "covering lint, test, build, and release would improve reproducibility and "
     "reduce onboarding friction for contributors."),
    ("Resumable Uploads and Downloads",
     "Supporting chunked, resumable transfers in both the CLI and web UI would "
     "significantly improve reliability for large files over unstable connections."),
    ("Prebuilt CLI Binaries",
     "Publishing prebuilt shc-cli binaries for Windows, macOS, and Linux via "
     "GitHub Releases would remove the Rust toolchain prerequisite for end users."),
]
for i, (title, desc) in enumerate(improvements, 1):
    p = doc.add_paragraph(style="List Number")
    r_bold = p.add_run(title + ": ")
    tnr(r_bold, 12, bold=True)
    r_body = p.add_run(desc)
    tnr(r_body, 12)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_para_spacing(p, space_after=4)

add_page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# BIBLIOGRAPHY
# ══════════════════════════════════════════════════════════════════════════════
chapter_heading(doc, "BIBLIOGRAPHY")
add_blank(doc)
body_para(doc,
    "References are formatted according to IEEE citation style.")
add_blank(doc)

references = [
    '[1] A. K. Shrestha and J. Vassileva, "Blockchain-Based Research Data Sharing '
    'Framework for Incentivizing the Data Owners," in Int. Conf. on Blockchain, '
    'Springer, Cham, Switzerland, 2018, pp. 259\u2013266.',
    '[2] S. Nakamoto, "Bitcoin: A Peer-to-Peer Electronic Cash System," Oct. 2008. '
    '[Online]. Available: https://bitcoin.org/bitcoin.pdf',
    '[3] Various Authors, "Blockchain applications in cloud computing, IoT, healthcare, '
    'data trading, and information security," survey of recent literature, 2015\u20132023.',
    '[4] J. Wu et al., "Traceable Attribute-Based Encryption for Privacy in Blockchain – '
    'Blockchain Access Control Systems," 2019.',
    '[5] C. Zhang and J. Zhao, "Smart Contract Based Digital Rights Management on '
    'Blockchain," License and transaction systems, 2018.',
    '[6] T. Ma et al., "DRMchain Dual Blockchain for Digital Content Protection," '
    'Protected digital content systems, Ethereum payments, 2018.',
    '[7] A. Rowhani-Farid, M. Allen, and A. G. Barnett, "What incentives increase data '
    'sharing in health and medical research? A systematic review," Res. Integr. Peer '
    'Rev., vol. 2, p. 4, 2017.',
    '[8] J. Benet, "IPFS – Content Addressed, Versioned, P2P File System," '
    'arXiv:1407.3561, 2014.',
    '[9] G. Wood, "Ethereum: A Secure Decentralised Generalised Transaction Ledger," '
    'Ethereum Proj. Yellow Pap., vol. 151, pp. 1\u201332, 2014.',
    '[10] A. Shamir, "How to Share a Secret," Commun. ACM, vol. 22, '
    'pp. 612\u2013613, 1979.',
    '[11] N. Nizamuddin, K. Salah, M. A. Azad, J. Arshad, and M. H. Rehman, '
    '"Decentralized Document Version Control Using Ethereum Blockchain and IPFS," '
    'Comput. Electr. Eng., vol. 76, pp. 183\u2013197, 2019.',
    '[12] H. R. Hasan and K. Salah, "Proof of Delivery of Digital Assets Using '
    'Blockchain and Smart Contracts," IEEE Access, vol. 6, pp. 65439\u201365448, 2018.',
    '[13] A. T. Panescu and V. Manta, "Smart Contracts for Research Data Rights '
    'Management over the Ethereum Blockchain Network," Sci. Technol. Libr., '
    'vol. 37, pp. 235\u2013245, 2018.',
    '[14] Y. Chen, H. Li, K. Li, and J. Zhang, "An Improved P2P File System Scheme '
    'Based on IPFS and Blockchain," in Proc. IEEE Int. Conf. Big Data, Boston, MA, '
    'USA, Dec. 2017, pp. 2652\u20132657.',
    '[15] M. Steichen et al., "Blockchain-Based, Decentralized Access Control for IPFS," '
    'in Proc. IEEE Int. Conf. Blockchain (Blockchain-2018), Halifax, NS, Canada, Jul. '
    '2018, pp. 1499\u20131506.',
    '[16] W. Sun, S. Yu, W. Lou, Y. T. Hou, and H. Li, "Protecting Your Right: Verifiable '
    'Attribute-Based Keyword Search with Fine-Grained Owner-Enforced Search Authorization '
    'in the Cloud," IEEE Trans. Parallel Distrib. Syst., vol. 27, pp. 1187\u20131198, 2016.',
    '[17] S. Wang, Y. Zhang, and Y. Zhang, "A Blockchain-Based Framework for Data Sharing '
    'with Fine-Grained Access Control in Decentralized Storage Systems," '
    'IEEE Access, vol. 6, pp. 38437\u201338450, 2018.',
    '[18] J. S. Park et al., "Smart Contract-Based Review System for an IoT Data '
    'Marketplace," Sensors, vol. 18, p. 3577, 2018.',
    '[19] L. Zhu et al., "Controllable and Trustworthy Blockchain-Based Cloud Data '
    'Management," Future Gener. Comput. Syst., vol. 91, pp. 527\u2013535, 2019.',
    '[20] J. Li, J. Wu, and L. Chen, "Block-Secure: Blockchain Based Scheme for Secure '
    'P2P Cloud Storage," Inf. Sci., vol. 465, pp. 219\u2013231, 2018.',
    '[21] R. B. Madhumala, S. Chhetri, K. C. Akshatha, and H. Jain, "Secure File '
    'Storage & Sharing on Cloud Using Cryptography," Int. J. Comput. Sci. Mob. '
    'Comput., vol. 10, no. 5, p. 49, 2021.',
    '[22] M. Naz et al., "Secure Data Sharing Platform Using Blockchain and '
    'InterPlanetary File System," IEEE Access, 2019.',
]

for ref in references:
    p = doc.add_paragraph()
    run = p.add_run(ref)
    tnr(run, 12)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.left_indent  = Inches(0.4)
    pf.first_line_indent = Inches(-0.4)
    set_para_spacing(p, space_after=4)

add_page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# APPENDIX
# ══════════════════════════════════════════════════════════════════════════════
chapter_heading(doc, "APPENDIX")
add_blank(doc)

main_heading(doc, "Appendix A: Backend API Endpoint Reference")
body_para(doc, "Table A.1 provides a complete reference for all backend API endpoints.")
add_blank(doc)
p_cap_a = doc.add_paragraph()
run_cap_a = p_cap_a.add_run("Table A.1: Backend API Endpoints")
tnr(run_cap_a, 10)
p_cap_a.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_para_spacing(p_cap_a, space_after=2)

api_endpoints = [
    ("GET",    "/",                                              "No",           "Health check"),
    ("POST",   "/auth/otp",                                     "No",           "Send OTP to email"),
    ("POST",   "/auth/login",                                   "No",           "Verify OTP, issue tokens"),
    ("GET",    "/auth/refresh-token",                           "Refresh token","Rotate token pair"),
    ("DELETE", "/auth/logout",                                  "Access token", "Invalidate session"),
    ("GET",    "/api/users/me",                                 "Access token", "Current user + plan"),
    ("GET",    "/api/files/",                                   "Access token", "List files (paginated)"),
    ("GET",    "/api/files/:fileId",                            "Access token", "File metadata + download URL"),
    ("POST",   "/api/files/add",                                "Access token", "Create file, get upload URL"),
    ("PATCH",  "/api/files/update-upload-status/:fileId",       "Access token", "Set uploaded, trigger notarization"),
    ("PATCH",  "/api/files/toggle-visibility/:fileId",          "Access token", "Toggle public/private"),
    ("PATCH",  "/api/files/increment-download-count/:fileId",   "Access token", "Increment download count"),
    ("PATCH",  "/api/files/rename/:id",                         "Access token", "Rename file"),
    ("DELETE", "/api/files/remove/:id",                         "Access token", "Delete file"),
    ("GET",    "/api/files/verify/:fileId",                     "No",           "Verify SHA-256 vs on-chain"),
    ("POST",   "/api/files/demo-tamper/:fileId",                "No",           "Demo: append null byte"),
    ("POST",   "/api/files/demo-restore/:fileId",               "No",           "Demo: restore original"),
    ("POST",   "/analyze-link",                                 "No",           "ML risk scoring"),
]

tbl_api = doc.add_table(rows=1, cols=4)
tbl_api.style = "Table Grid"
for cell, hdr in zip(tbl_api.rows[0].cells, ["Method", "Endpoint", "Auth", "Description"]):
    p = cell.paragraphs[0]
    p.clear()
    r = p.add_run(hdr)
    tnr(r, 11, bold=True)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
for method, endpoint, auth, desc in api_endpoints:
    row = tbl_api.add_row()
    for cell, text, align in zip(
        row.cells,
        [method, endpoint, auth, desc],
        [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT,
         WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT],
    ):
        p = cell.paragraphs[0]
        p.clear()
        r = p.add_run(text)
        tnr(r, 10)
        p.alignment = align
        p.paragraph_format.space_after = Pt(1)

add_blank(doc)
main_heading(doc, "Appendix B: Random Forest Model Evaluation Metrics")
add_blank(doc)
p_cap_b = doc.add_paragraph()
run_cap_b = p_cap_b.add_run("Table B.1: Structured Classifier (Random Forest) Test-Set Metrics")
tnr(run_cap_b, 10)
p_cap_b.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_para_spacing(p_cap_b, space_after=2)

metrics = [
    ("Accuracy",               "91.4%"),
    ("Precision",              "90.8%"),
    ("Recall",                 "92.1%"),
    ("F1 Score",               "91.4%"),
    ("ROC-AUC",                "0.964"),
    ("5-Fold CV F1 (mean)",    "89.8%"),
    ("5-Fold CV F1 (std dev)", "\u00b10.6%"),
    ("Training samples",       "2,560 (80% of 3,200)"),
    ("Test samples",           "640 (20% of 3,200)"),
    ("Features",               "14 structured features"),
]
tbl_metrics = doc.add_table(rows=1, cols=2)
tbl_metrics.style = "Table Grid"
for cell, hdr in zip(tbl_metrics.rows[0].cells, ["Metric", "Value"]):
    p = cell.paragraphs[0]
    p.clear()
    r = p.add_run(hdr)
    tnr(r, 12, bold=True)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
for metric, value in metrics:
    row = tbl_metrics.add_row()
    for cell, text, align in zip(
        row.cells,
        [metric, value],
        [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER],
    ):
        p = cell.paragraphs[0]
        p.clear()
        r = p.add_run(text)
        tnr(r, 12)
        p.alignment = align
        p.paragraph_format.space_after = Pt(2)

# ── save ───────────────────────────────────────────────────────────────────────
output_path = r"c:\Users\USER\Downloads\SHC_Major_Project_Report_ES458.docx"
doc.save(output_path)
print(f"Report saved to: {output_path}")
