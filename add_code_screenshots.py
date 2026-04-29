"""
Add 15 syntax-highlighted code screenshots to the SHC ES-458 report.
Each screenshot has a bold figure title directly below it, plus a
caption line with the file path. The List of Figures table inside
the document is also updated to include all 15 entries.
"""

from pathlib import Path
from io import BytesIO

from pygments import highlight
from pygments.lexers import GoLexer, RustLexer, PythonLexer
from pygments.formatters import ImageFormatter
from pygments.styles import get_style_by_name

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ── code snippets (15 total) ───────────────────────────────────────────────────

SNIPPETS = [
    # ── 1 ──────────────────────────────────────────────────────────────────────
    {
        "fig_no":  "Figure 3.1",
        "title":   "BlockchainService — Struct Definition and Initialization (Go)",
        "caption": "shc-backend/services/blockchain.service.go",
        "lexer":   GoLexer(),
        "code": """\
// BlockchainService submits SHA-256 file hashes to Ethereum as on-chain
// notarizations via a 0-ETH self-send whose `data` field contains
// "shc:" + hex(sha256_bytes) — pure-Go, no CGo required.
type BlockchainService struct {
    rpcURL     string
    privateKey *secp256k1.PrivateKey
    address    string   // "0x..." derived from private key
    chainID    *big.Int
    enabled    bool
}

func NewBlockchainService() *BlockchainService {
    rpcURL := strings.TrimSpace(os.Getenv("ETH_RPC_URL"))
    rawKey := strings.TrimSpace(os.Getenv("ETH_WALLET_PRIVATE_KEY"))
    if rpcURL == "" || rawKey == "" {
        log.Println("BlockchainService: notarization disabled")
        return &BlockchainService{enabled: false}
    }
    keyBytes, _ := hex.DecodeString(strings.TrimPrefix(rawKey, "0x"))
    privKey      := secp256k1.PrivKeyFromBytes(keyBytes)
    address      := ethAddressFromPrivKey(privKey)
    chainID      := big.NewInt(11155111) // Sepolia testnet default
    if raw := strings.TrimSpace(os.Getenv("ETH_CHAIN_ID")); raw != "" {
        id := new(big.Int)
        if _, ok := id.SetString(raw, 10); ok { chainID = id }
    }
    return &BlockchainService{
        rpcURL: rpcURL, privateKey: privKey,
        address: address, chainID: chainID, enabled: true,
    }
}""",
    },
    # ── 2 ──────────────────────────────────────────────────────────────────────
    {
        "fig_no":  "Figure 3.2",
        "title":   "EIP-155 Transaction Signing — NotarizeHash and buildSignedTx (Go)",
        "caption": "shc-backend/services/blockchain.service.go",
        "lexer":   GoLexer(),
        "code": """\
func (bs *BlockchainService) NotarizeHash(
    ctx context.Context, sha256Hash []byte,
) (string, error) {
    nonce,    _ := bs.pendingNonce(ctx)
    gasPrice, _ := bs.suggestGasPrice(ctx)
    data        := []byte("shc:" + hex.EncodeToString(sha256Hash))
    gasLimit    := uint64(21000 + 68*uint64(len(data)))
    rawTx       := bs.buildSignedTx(nonce, gasPrice, gasLimit, data)
    return bs.sendRawTransaction(ctx, rawTx)
}

// buildSignedTx assembles and EIP-155-signs a legacy Ethereum transaction.
func (bs *BlockchainService) buildSignedTx(
    nonce uint64, gasPrice *big.Int, gasLimit uint64, data []byte,
) []byte {
    toBytes, _ := hex.DecodeString(strings.TrimPrefix(bs.address, "0x"))
    // EIP-155 pre-image includes chainID to prevent replay on other networks
    preimage := rlpList(
        rlpUint(nonce), rlpBigInt(gasPrice), rlpUint(gasLimit),
        rlpBytes(toBytes), rlpBigInt(big.NewInt(0)), rlpBytes(data),
        rlpBigInt(bs.chainID), rlpBytes(nil), rlpBytes(nil),
    )
    msgHash  := ethKeccak256(preimage)
    compact  := decredecdsa.SignCompact(bs.privateKey, msgHash, false)
    recovID  := int64(compact[0]) - 27
    rInt     := new(big.Int).SetBytes(compact[1:33])
    sInt     := new(big.Int).SetBytes(compact[33:65])
    v := new(big.Int).SetInt64(recovID + 35)
    v.Add(v, new(big.Int).Mul(bs.chainID, big.NewInt(2)))
    return rlpList(
        rlpUint(nonce), rlpBigInt(gasPrice), rlpUint(gasLimit),
        rlpBytes(toBytes), rlpBigInt(big.NewInt(0)), rlpBytes(data),
        rlpBigInt(v), rlpBigInt(rInt), rlpBigInt(sInt),
    )
}""",
    },
    # ── 3 ──────────────────────────────────────────────────────────────────────
    {
        "fig_no":  "Figure 3.3",
        "title":   "JWT Authentication Middleware with Public-Route Bypass (Go)",
        "caption": "shc-backend/middlewares/auth.middleware.go",
        "lexer":   GoLexer(),
        "code": """\
var publicVerifyRoute    = regexp.MustCompile(`^/api/files/verify/[0-9a-fA-F-]+$`)
var publicDemoTamperRoute = regexp.MustCompile(`^/api/files/demo-tamper/[0-9a-fA-F-]+$`)
var publicDemoRestoreRoute = regexp.MustCompile(`^/api/files/demo-restore/[0-9a-fA-F-]+$`)

func AuthMiddleware(c fiber.Ctx, as *services.AppService) error {
    accessToken := string(c.Request().Header.Peek("Authorization"))
    if accessToken == "" {
        accessToken = c.Cookies("__shc_access_token")
    }
    accessToken = strings.TrimPrefix(accessToken, "Bearer ")

    claim, err := as.AuthService.VerifyAccessToken(accessToken)
    semiPublic  := isSemiPublicRoute(c.Method(), c.Path())

    if err == nil {
        c.Request().Header.Set("user_id",    claim.ID.String())
        c.Request().Header.Set("user_email", claim.Email)
        c.Request().Header.Set("user_name",  claim.Name)
        c.Request().Header.Set("token_type", claim.TokenType)
    }
    if !semiPublic && err != nil {
        return c.SendStatus(401)
    }
    return c.Next()
}""",
    },
    # ── 4 ──────────────────────────────────────────────────────────────────────
    {
        "fig_no":  "Figure 3.4",
        "title":   "OTP Verification and JWT Token Issuance — Login Handler (Go)",
        "caption": "shc-backend/handlers/auth-handlers/login.handler.go",
        "lexer":   GoLexer(),
        "code": """\
func VerifyOtpAndGetTokens(c fiber.Ctx, as *services.AppService) error {
    req := new(CheckOtpRequest)
    if err := c.Bind().Body(req); err != nil {
        return err
    }
    otp, err := strconv.Atoi(req.Otp)
    if err != nil {
        return &fiber.Error{Code: fiber.StatusBadRequest, Message: "Invalid OTP"}
    }
    if err = as.AuthService.VerifyOtp(req.Email, otp); err != nil {
        return &fiber.Error{Code: fiber.StatusUnauthorized, Message: err.Error()}
    }

    var user *m.User
    u, err := as.UserService.FindUserByEmail(req.Email)
    if err != nil {
        // First login: create new user record
        user, err = as.UserService.CreateUser(
            &m.User{Email: req.Email, Name: req.Name})
    } else {
        // Returning user: update name if changed
        user, err = as.UserService.UpdateAUser(
            &m.User{ID: u.ID, Email: req.Email, Name: req.Name})
    }
    if err != nil { return err }

    tokens, err := as.AuthService.GenerateTokens(user.ID, user.Name, user.Email)
    if err != nil { return err }

    return c.JSON(fiber.Map{
        "access_token":  tokens.AccessToken,
        "refresh_token": tokens.RefreshToken,
        "name":          user.Name,
        "email":         user.Email,
        "id":            user.ID,
    })
}""",
    },
    # ── 5 ──────────────────────────────────────────────────────────────────────
    {
        "fig_no":  "Figure 3.5",
        "title":   "JWT Token Generation — Access and Refresh Token Claims (Go)",
        "caption": "shc-backend/services/auth.service.go",
        "lexer":   GoLexer(),
        "code": """\
type AccessTokenClaim struct {
    jwt.RegisteredClaims
    ID        uuid.UUID
    Email     string
    Name      string
    TokenType string
}

func (a *AuthService) generateAccessToken(
    id uuid.UUID, name, email string,
) (string, error) {
    token := jwt.NewWithClaims(jwt.SigningMethodHS256,
        AccessTokenClaim{
            ID: id, Name: name, Email: email,
            TokenType: "access.token",
            RegisteredClaims: jwt.RegisteredClaims{
                ExpiresAt: jwt.NewNumericDate(
                    time.Now().Add(a.JwtAccessTokenExpiresIn)),
                Issuer:   "shc.auth.service",
                IssuedAt: jwt.NewNumericDate(time.Now()),
            },
        })
    return token.SignedString([]byte(a.JwtSecretKey))
}

func (a *AuthService) GenerateTokens(
    id uuid.UUID, name, email string,
) (*Tokens, error) {
    access,  err := a.generateAccessToken(id, name, email)
    if err != nil { return nil, err }
    refresh, err := a.generateRefreshToken(id, name, email)
    if err != nil { return nil, err }
    return &Tokens{AccessToken: access, RefreshToken: refresh}, nil
}""",
    },
    # ── 6 ──────────────────────────────────────────────────────────────────────
    {
        "fig_no":  "Figure 3.6",
        "title":   "File Upload Initialization — Presigned PUT URL Generation (Go)",
        "caption": "shc-backend/handlers/file-handlers/add-file.handler.go",
        "lexer":   GoLexer(),
        "code": """\
func AddFileToDb(c fiber.Ctx, as *services.AppService) error {
    userId, _ := uuid.Parse(
        string(c.Request().Header.Peek("user_id")))

    body := new(AddFileDto)
    if err := c.Bind().Body(body); err != nil {
        return fiber.NewError(fiber.StatusBadRequest, "Invalid request body")
    }

    // Normalize MIME: browsers report non-standard types for executables
    mimeType := strings.TrimSpace(body.MimeType)
    if mimeType == "" { mimeType = "application/octet-stream" }
    if mimeType == "application/vnd.microsoft.portable-executable" {
        mimeType = "application/x-msdownload"
    }

    key := aws.String(userId.String() + "/" +
        uuid.NewString() + "_" +
        strings.Replace(body.FileName, " ", "_", -1))

    res, err := as.S3Service.S3PresignClient.PresignPutObject(
        context.Background(),
        &s3.PutObjectInput{
            Bucket: aws.String(as.S3Service.BucketName),
            Key:    key,
        },
        s3.WithPresignExpires(15*time.Minute),
    )
    if err != nil { return err }

    // ... quota check, DB insert, return presigned URL + fileId
    return c.JSON(fiber.Map{
        "upload_url": res.URL,
        "file_id":    newFile.ID,
    })
}""",
    },
    # ── 7 ──────────────────────────────────────────────────────────────────────
    {
        "fig_no":  "Figure 3.7",
        "title":   "File Retrieval — Presigned GET URL with Redis Cache (Go)",
        "caption": "shc-backend/handlers/file-handlers/get-file.handler.go",
        "lexer":   GoLexer(),
        "code": """\
func GetFile(c fiber.Ctx, as *services.AppService) error {
    fileId, _ := uuid.Parse(c.Params("fileId"))
    userId, _ := uuid.Parse(
        string(c.Request().Header.Peek("user_id")))

    file, err := as.FileService.FindFileById(fileId)
    if err != nil { return err }

    if file.UserId != userId && !file.IsPublic {
        return &fiber.Error{Code: fiber.StatusForbidden,
            Message: "You are not allowed to access this file"}
    }
    // 410 Gone after 48-hour expiry
    if file.ExpiresAt != nil && file.ExpiresAt.Before(time.Now()) {
        return &fiber.Error{Code: fiber.StatusGone,
            Message: "This file has expired and is no longer available"}
    }

    as.FileService.IncrementViewCount(fileId)

    // Check Redis cache before generating a new presigned URL
    downloadUrl := ""
    cacheKey    := "download_url_of_" + file.R2Path
    if cached, err := as.RedisService.GetCache(cacheKey); err != nil {
        res, _ := as.S3Service.S3PresignClient.PresignGetObject(
            c.Context(),
            &s3.GetObjectInput{
                Bucket: aws.String(as.S3Service.BucketName),
                Key:    aws.String(file.R2Path),
            },
        )
        as.RedisService.SetCache(cacheKey, res.URL, 880*time.Second)
        downloadUrl = res.URL
    } else {
        downloadUrl, _ = cached.(string)
    }
    // Risk score + file metadata returned together
    return c.JSON(buildFileResponse(file, downloadUrl, risk))
}""",
    },
    # ── 8 ──────────────────────────────────────────────────────────────────────
    {
        "fig_no":  "Figure 3.8",
        "title":   "File GORM Model — Notarization and Integrity Fields (Go)",
        "caption": "shc-backend/models/file.go",
        "lexer":   GoLexer(),
        "code": """\
// IntegrityStatus reflects the result of blockchain-based integrity check.
// "unverified" — file has not been notarized or verification not run.
// "verified"   — on-chain hash matches current file bytes (tamper-proof).
// "tampered"   — on-chain hash does NOT match; content has changed.
type IntegrityStatus string

const (
    IntegrityUnverified IntegrityStatus = "unverified"
    IntegrityVerified   IntegrityStatus = "verified"
    IntegrityTampered   IntegrityStatus = "tampered"
)

type File struct {
    gorm.Model
    ID              uuid.UUID       `gorm:"type:uuid;default:gen_random_uuid()" json:"id"`
    UploadStatus    UploadStatus    `gorm:"size:255;not null;default:not_started" json:"upload_status"`
    Name            string          `gorm:"size:255;not null" json:"name"`
    Size            uint            `gorm:"not null" json:"size"`
    IsPublic        bool            `gorm:"not null;default:false" json:"is_public"`
    MimeType        string          `gorm:"size:255;not null" json:"mime_type"`
    R2Path          string          `gorm:"size:255;not null" json:"r2_path"`
    DownloadCount   uint            `gorm:"not null;default:0" json:"download_count"`
    UserId          uuid.UUID       `gorm:"not null;index" json:"user_id"`
    ExpiresAt       *time.Time      `gorm:"index" json:"expires_at"`
    NotarizationTx  string          `gorm:"size:255;default:''" json:"notarization_tx"`
    IntegrityStatus IntegrityStatus `gorm:"size:32;not null;default:'unverified'" json:"integrity_status"`
    SHA256Hash      string          `gorm:"size:64;default:''" json:"sha256_hash"`
}""",
    },
    # ── 9 ──────────────────────────────────────────────────────────────────────
    {
        "fig_no":  "Figure 3.9",
        "title":   "Cron Jobs — Expired File Cleanup and Startup Backfill (Go)",
        "caption": "shc-backend/cmd/jobs.go",
        "lexer":   GoLexer(),
        "code": """\
func runCronJobs(as *services.AppService) {
    // Delete non-uploaded stale records at midnight
    as.CronService.AddFunc("@midnight", func() {
        as.FileService.DeleteAllNonUploadedFiles()
    })

    // Purge expired files from storage every hour
    as.CronService.AddFunc("@every 1h", func() {
        r2Paths, err := as.FileService.DeleteExpiredFiles()
        if err != nil {
            log.Printf("DeleteExpiredFiles: %v", err)
            return
        }
        for _, key := range r2Paths {
            if err := as.S3Service.DeleteObject(key); err != nil {
                log.Printf("S3 delete failed for %q: %v", key, err)
            }
        }
        log.Printf("Deleted %d expired file(s) from storage", len(r2Paths))
    })
    as.CronService.Start()
}

// backfillIntegrityHashes runs once at startup: finds uploaded files
// without a SHA-256 hash, computes + stores it, sets status="verified".
func backfillIntegrityHashes(as *services.AppService) {
    files, _ := as.FileService.FindUploadedFilesWithoutHash()
    log.Printf("backfill: computing hashes for %d file(s)...", len(files))
    for _, file := range files {
        ctx, cancel := context.WithTimeout(context.Background(), 2*time.Minute)
        defer cancel()
        fileBytes, _ := fetchBytesForBackfill(ctx, as, &file)
        sum := sha256.Sum256(fileBytes)
        as.FileService.SetSHA256Hash(file.ID, hex.EncodeToString(sum[:]))
    }
}""",
    },
    # ── 10 ─────────────────────────────────────────────────────────────────────
    {
        "fig_no":  "Figure 3.10",
        "title":   "Paginated File Listing with Redis Cache (Go)",
        "caption": "shc-backend/handlers/file-handlers/list-files.handler.go",
        "lexer":   GoLexer(),
        "code": """\
func ListFiles(c fiber.Ctx, as *services.AppService) error {
    userIdString := string(c.Request().Header.Peek("user_id"))
    page,  _ := strconv.Atoi(c.Query("page"))
    limit, _ := strconv.Atoi(c.Query("limit"))
    if page  < 1 { page  = 1  }
    if limit < 1 { limit = 10 }

    search   := c.Query("search")
    language := c.Query("language")

    userId, err := uuid.Parse(userIdString)
    if err != nil { return err }

    cacheKey := services.UserFilesCacheKey(userId, search, language, page, limit)
    var cachedResponse listFilesResponse
    // Serve from Redis if available — avoids DB round-trip on repeat loads
    if err := as.RedisService.GetJSONCache(cacheKey, &cachedResponse); err == nil {
        return c.JSON(cachedResponse)
    }

    filesPaginationResults, err := as.FileService.FindFilesByUserId(
        userId, search, language, page, limit)
    if err != nil { return err }

    response := buildListFilesResponse(filesPaginationResults)
    as.RedisService.SetJSONCache(cacheKey, response, 30*time.Second)
    return c.JSON(response)
}""",
    },
    # ── 11 ─────────────────────────────────────────────────────────────────────
    {
        "fig_no":  "Figure 3.11",
        "title":   "Rust CLI — File Upload with Progress Bar (Rust)",
        "caption": "shc-cli/src/command/add.rs",
        "lexer":   RustLexer(),
        "code": """\
pub async fn upload_file(
    file_path: &Path,
    api_client: &mut api_client::ApiClient,
) -> Result<(), Box<dyn std::error::Error>> {
    // Auto-zip directories before upload
    let file_path = if file_path.is_dir() {
        let pb = ProgressBar::new_spinner();
        pb.enable_steady_tick(Duration::from_millis(200));
        pb.set_message("Compressing folder...");
        let zipped = zip_directory_recursive(file_path, 30 * 1024 * 1024)?;
        pb.finish_and_clear();
        zipped
    } else {
        file_path.to_path_buf()
    };

    let file_name  = file_path.file_name().unwrap().to_str().unwrap();
    let mime_type  = mime_guess::from_path(&file_path).first_or_octet_stream();
    let file       = tokio::fs::File::open(&file_path).await?;
    let total_size = file.metadata().await?.len();

    // Stream upload with live progress bar
    let pb = ProgressBar::new(total_size);
    pb.set_style(ProgressStyle::with_template(
        "{spinner:.green} [{bar:40.cyan/blue}] {bytes}/{total_bytes} ({eta})"
    ).unwrap().progress_chars("=>-"));

    let stream = ReaderStream::new(file).map(move |chunk| {
        chunk.inspect(|c| pb.inc(c.len() as u64))
    });
    let body = reqwest::Body::wrap_stream(stream);
    // PUT directly to presigned R2 URL — bypasses the backend server
    api_client.upload_to_presigned_url(&upload_url, body, total_size).await?;
    Ok(())
}""",
    },
    # ── 12 ─────────────────────────────────────────────────────────────────────
    {
        "fig_no":  "Figure 3.12",
        "title":   "Rust CLI — File Listing with TUI Table Output (Rust)",
        "caption": "shc-cli/src/command/list.rs",
        "lexer":   RustLexer(),
        "code": """\
pub async fn list_files(
    search: &str,
    api_client: &mut api_client::ApiClient,
) -> Result<(), Box<dyn std::error::Error>> {
    let pb = ProgressBar::new_spinner();
    pb.enable_steady_tick(Duration::from_millis(200));
    pb.set_message("Fetching files...");

    let res = api_client.list_files(search).await?;
    pb.finish_and_clear();

    if res.results.is_empty() {
        println!("No files found.");
        return Ok(());
    }

    let prompt = format!("Select a file ({} total).", res.results.len());
    let selection = shc_file_input(&res.results, &prompt);
    let file = &res.results[selection];

    let size_str = if file.size < 1024 {
        format!("{:.3} KB", file.size as f64 / 1024.0)
    } else {
        format!("{:.3} MB", file.size as f64 / 1024.0 / 1024.0)
    };
    let share_link = format!(
        "https://shc-frontend-two.vercel.app/share/{}", file.id);

    console::Term::stdout().write_line(&format!(
        "\\nName: {}\\nStatus: {}\\nSize: {}\\nVisibility: {}\\nShare: {}",
        style(&file.name).cyan(),
        style(&file.upload_status).yellow(),
        style(size_str).magenta(),
        style(if file.is_public {"Public"} else {"Private"}).blue(),
        style(&share_link).underlined().bright().blue(),
    ))?;
    Ok(())
}""",
    },
    # ── 13 ─────────────────────────────────────────────────────────────────────
    {
        "fig_no":  "Figure 3.13",
        "title":   "Rust CLI — Transparent JWT Refresh in ApiClient (Rust)",
        "caption": "shc-cli/src/api_client.rs",
        "lexer":   RustLexer(),
        "code": """\
impl ApiClient {
    // Intercepts 401 responses, refreshes token, retries automatically.
    async fn refresh_token(
        &mut self,
    ) -> Result<(), Box<dyn std::error::Error>> {
        if self.tried_refreshing_token {
            self.user_config.clear();
            self.login_again();
        }
        let res = self.client
            .get(format!("{}/auth/refresh-token", self.api_base_url))
            .header(
                "Authorization",
                self.user_config.user.refresh_token.as_ref().unwrap(),
            )
            .send().await?;

        match res.status() {
            reqwest::StatusCode::OK => {
                let data = res.json::<RefreshTokenResponse>().await?;
                self.user_config.user.access_token  = Some(data.access_token);
                self.user_config.user.refresh_token = Some(data.refresh_token);
                self.user_config.save(); // persist new tokens to disk
                self.tried_refreshing_token = true;
                Ok(())
            }
            _ => { self.login_again(); Ok(()) }
        }
    }
}""",
    },
    # ── 14 ─────────────────────────────────────────────────────────────────────
    {
        "fig_no":  "Figure 3.14",
        "title":   "Hybrid Risk Scoring — Rule Engine with Blockchain Override (Python)",
        "caption": "shc-risk-ml-service/app/rules.py",
        "lexer":   PythonLexer(),
        "code": """\
def baseline_score(
    features: Dict[str, float], evidence: List[str],
) -> Tuple[int, List[str]]:
    score = 0
    reasons = list(evidence)

    if features.get("is_executable_ext", 0):   score += 28
    if features.get("contains_macro_indicator", 0): score += 24
    if features.get("domain_risk", 0) > 0.55:  score += 18
    if features.get("size_anomaly", 0):         score += 10
    if features.get("entropy", 0) > 5.0:        score += 20
    elif features.get("entropy", 0) > 4.4:      score += 12
    if features.get("known_bad_hash", 0):       score += 45

    keyword_hits = int(features.get("keyword_hit_count", 0))
    if keyword_hits > 0:
        score += min(18, keyword_hits * 4)
        reasons.append("Contains social engineering keywords")

    # Blockchain integrity is a hard override — it overrules all ML signals
    blockchain = features.get("blockchain_integrity", 0.0)
    if blockchain > 0:
        score += 55
        reasons.append(
            "Blockchain integrity FAILED: file content does not match "
            "the on-chain hash (tampering detected)")
    elif blockchain < 0:
        score = max(0, score - 25)
        reasons.append(
            "Blockchain integrity verified: file matches its "
            "immutable on-chain hash")

    return max(0, min(100, score)), _dedupe(reasons)[:6]""",
    },
    # ── 15 ─────────────────────────────────────────────────────────────────────
    {
        "fig_no":  "Figure 3.15",
        "title":   "Model Training — Random Forest + TF-IDF Pipeline (Python)",
        "caption": "shc-risk-ml-service/training/train_models.py",
        "lexer":   PythonLexer(),
        "code": """\
FEATURE_ORDER = [
    "is_executable_ext", "is_archive", "contains_macro_indicator",
    "domain_risk", "entropy", "keyword_hit_count", "size_mb",
    "size_anomaly", "share_frequency", "download_frequency",
    "unknown_upload_source", "known_bad_hash", "text_length",
    "blockchain_integrity",  # +1 tampered / 0 unverified / -1 verified
]

def main() -> None:
    X, y = _generate_structured_dataset(seed=42, size=3200)
    X_train, X_test, y_train, y_test = train_test_split(
        X, y, test_size=0.2, random_state=42, stratify=y)

    # Structured model: Random Forest (220 trees, max_depth=10)
    rf = RandomForestClassifier(
        n_estimators=220, max_depth=10,
        random_state=42, class_weight="balanced")
    rf.fit(X_train, y_train)

    y_pred = rf.predict(X_test)
    y_prob = rf.predict_proba(X_test)[:, 1]
    cv     = cross_val_score(rf, X, y, cv=5, scoring="f1")
    print(f"Accuracy: {accuracy_score(y_test, y_pred):.4f}")
    print(f"ROC-AUC:  {roc_auc_score(y_test, y_prob):.4f}")
    print(f"CV F1:    {cv.mean():.4f} ± {cv.std():.4f}")

    # Text model: TF-IDF + Logistic Regression pipeline
    text_model = Pipeline([
        ("tfidf", TfidfVectorizer(ngram_range=(1, 2), max_features=5000)),
        ("clf",   LogisticRegression(max_iter=1000, class_weight="balanced")),
    ])
    text_model.fit(X_text_train, y_text_train)
    joblib.dump(rf,         MODEL_DIR / "structured_model.joblib")
    joblib.dump(text_model, MODEL_DIR / "text_model.joblib")""",
    },
]

# ── image generation ───────────────────────────────────────────────────────────

STYLE_NAME = "monokai"
FONT_SIZE  = 13


def code_to_png_bytes(code: str, lexer, style_name: str = STYLE_NAME) -> bytes:
    formatter = ImageFormatter(
        style=get_style_by_name(style_name),
        font_size=FONT_SIZE,
        line_numbers=True,
        line_number_bg="#272822",
        line_number_fg="#888888",
        line_pad=3,
        image_pad=10,
    )
    return highlight(code, lexer, formatter)


# ── docx helpers ───────────────────────────────────────────────────────────────

from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph as DocxParagraph


def insert_paragraph_after(doc, ref_para):
    """Splice a new empty paragraph immediately after ref_para."""
    new_p = OxmlElement('w:p')
    ref_para._p.addnext(new_p)
    return DocxParagraph(new_p, doc)


def tnr(run, size_pt, bold=False, italic=False):
    run.font.name   = "Times New Roman"
    run.font.size   = Pt(size_pt)
    run.font.bold   = bold
    run.font.italic = italic


def insert_code_figure(doc, ref_para, snippet, png_bytes, width_inches=5.8):
    """
    Insert after ref_para:
      1. Image (centered, space_before=10)
      2. Bold title line  — "Figure 3.X: <title>"   (10 pt, centered)
      3. Italic file path — "<path>"                 (10 pt, centered)
    Returns the last inserted paragraph.
    """
    stream = BytesIO(png_bytes)

    # 1 – image
    img_para = insert_paragraph_after(doc, ref_para)
    img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    img_para.paragraph_format.space_before = Pt(10)
    img_para.paragraph_format.space_after  = Pt(2)
    img_para.add_run().add_picture(stream, width=Inches(width_inches))

    # 2 – bold title  "Figure 3.X: <title>"
    title_para = insert_paragraph_after(doc, img_para)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.paragraph_format.space_after = Pt(1)
    r = title_para.add_run(f"{snippet['fig_no']}: {snippet['title']}")
    tnr(r, 10, bold=True)

    # 3 – italic file path
    path_para = insert_paragraph_after(doc, title_para)
    path_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    path_para.paragraph_format.space_after = Pt(14)
    r2 = path_para.add_run(snippet["caption"])
    tnr(r2, 10, italic=True)

    return path_para


# ── load document ──────────────────────────────────────────────────────────────

# Always read from the original (v1); write to v2 to avoid clobbering.
REPORT_PATH = Path(r"c:\Users\USER\Downloads\SHC_Major_Project_Report_ES458.docx")
doc = Document(str(REPORT_PATH))

# ── Step 1: update List of Figures table ──────────────────────────────────────

LOF_HEADING = "LIST OF FIGURES"
lof_table = None
found_lof = False
for elem in doc.element.body:
    tag = elem.tag.split('}')[-1]
    if tag == 'p':
        from docx.text.paragraph import Paragraph as P
        p = P(elem, doc)
        if LOF_HEADING in p.text.upper():
            found_lof = True
    if found_lof and tag == 'tbl':
        from docx.table import Table
        lof_table = Table(elem, doc)
        break

if lof_table:
    print(f"Found List of Figures table ({len(lof_table.rows)} existing rows). Replacing entries...")
    # Remove all rows except the header
    for row in lof_table.rows[1:]:
        row._tr.getparent().remove(row._tr)

    # Add one row per snippet
    for s in SNIPPETS:
        row = lof_table.add_row()
        data = [
            (s["fig_no"],  WD_ALIGN_PARAGRAPH.CENTER),
            (s["title"],   WD_ALIGN_PARAGRAPH.LEFT),
            ("—",          WD_ALIGN_PARAGRAPH.CENTER),
        ]
        for cell, (text, align) in zip(row.cells, data):
            p = cell.paragraphs[0]
            p.clear()
            r = p.add_run(text)
            tnr(r, 12)
            p.alignment = align
            p.paragraph_format.space_after = Pt(2)
    print("List of Figures updated.")
else:
    print("Warning: List of Figures table not found — skipping update.")

# ── Step 2: find insertion point for code figures ─────────────────────────────

INSERT_AFTER_TEXT = "3.3 Technical Implementation"
target_idx = None
for i, para in enumerate(doc.paragraphs):
    if INSERT_AFTER_TEXT in para.text:
        target_idx = i
        break

if target_idx is None:
    target_idx = len(doc.paragraphs) - 1
    print(f"Warning: '{INSERT_AFTER_TEXT}' not found — appending near end.")
else:
    print(f"Insertion point: para {target_idx} → '{doc.paragraphs[target_idx].text[:60]}'")

# ── Step 3: insert intro + 15 code figures ────────────────────────────────────

ref_para = doc.paragraphs[target_idx]

intro_para = insert_paragraph_after(doc, ref_para)
intro_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
intro_para.paragraph_format.space_after = Pt(8)
r_intro = intro_para.add_run(
    "The following figures (Figures 3.1 – 3.15) present syntax-highlighted "
    "source code excerpts from each of the four SHC components, illustrating "
    "the key implementation details discussed in Sections 3.3.1 – 3.3.6."
)
tnr(r_intro, 12)

last_para = intro_para
print("\nGenerating code images...")

for snippet in SNIPPETS:
    print(f"  Rendering {snippet['fig_no']}: {snippet['title']}")
    png_bytes = code_to_png_bytes(snippet["code"], snippet["lexer"])
    last_para = insert_code_figure(doc, last_para, snippet, png_bytes)

# ── Step 4: save ──────────────────────────────────────────────────────────────

out_path = r"c:\Users\USER\Downloads\SHC_Major_Project_Report_ES458_v2.docx"
doc.save(out_path)
print(f"\nDone. Saved to: {out_path}")
